"""
Convert paper_manuscript.tex to Word (.docx) format.
Parses the LaTeX structure and renders it using python-docx.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEX_PATH = os.path.join(SCRIPT_DIR, 'paper_manuscript.tex')
OUT_PATH = os.path.join(SCRIPT_DIR, 'paper_manuscript.docx')


def clean_latex(text):
    """Remove common LaTeX markup from text."""
    # Handle \textbf, \textit, \emph (simplified - just extract content)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    # Citations
    text = re.sub(r'\\citep?\{([^}]*)\}', r'[\1]', text)
    # References
    text = re.sub(r'\\ref\{([^}]*)\}', r'\1', text)
    text = re.sub(r'Fig\.~\\ref\{([^}]*)\}', r'Fig. \1', text)
    text = re.sub(r'Table~\\ref\{([^}]*)\}', r'Table \1', text)
    # Math shortcuts
    text = re.sub(r'\$\\mathbf\{([^}]*)\}\$', r'\1', text)
    text = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathbb\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    # Inline math - simple cases
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    # Special chars
    text = text.replace('\\textdegree', '\u00b0')
    text = text.replace('\\%', '%')
    text = text.replace('\\&', '&')
    text = text.replace('\\$', '$')
    text = text.replace('\\#', '#')
    text = text.replace('\\,', ' ')
    text = text.replace('\\;', ' ')
    text = text.replace('\\!', '')
    text = text.replace('\\quad', '    ')
    text = text.replace('\\qquad', '        ')
    text = text.replace('~', ' ')
    text = text.replace('``', '\u201c')
    text = text.replace("''", '\u201d')
    text = text.replace('`', '\u2018')
    text = text.replace("'", '\u2019')
    text = text.replace('---', '\u2014')
    text = text.replace('--', '\u2013')
    text = text.replace('\\to', '\u2192')
    text = text.replace('\\times', '\u00d7')
    text = text.replace('\\pm', '\u00b1')
    text = text.replace('\\leq', '\u2264')
    text = text.replace('\\geq', '\u2265')
    text = text.replace('\\approx', '\u2248')
    text = text.replace('\\infty', '\u221e')
    text = text.replace('\\in', '\u2208')
    text = text.replace('\\ldots', '\u2026')
    text = text.replace('\\epsilon', '\u03b5')
    text = text.replace('\\lambda', '\u03bb')
    text = text.replace('\\gamma', '\u03b3')
    text = text.replace('\\beta', '\u03b2')
    text = text.replace('\\alpha', '\u03b1')
    text = text.replace('\\Delta', '\u0394')
    text = text.replace('\\bar{s}', 's\u0304')
    text = text.replace('\\sim', '~')
    text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)
    # Number formatting
    text = text.replace('{,}', ',')
    # Remove remaining simple commands
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    # Remove leftover braces
    text = text.replace('{', '').replace('}', '')
    # Clean up whitespace
    text = re.sub(r'  +', ' ', text)
    return text.strip()


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                             alignment=None, font_size=None, space_after=None,
                             space_before=None, first_line_indent=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean_latex(text))
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    # Set East Asian font
    rpr = run._element.get_or_add_rPr()
    ea_font = rpr.makeelement(qn('w:rFonts'), {})
    ea_font.set(qn('w:eastAsia'), 'SimSun')
    rpr.insert(0, ea_font)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_rich_paragraph(doc, text, style='Normal', space_after=6, space_before=0):
    """Add paragraph with bold/italic markup from LaTeX."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)

    # Split by bold/italic markers
    # First handle \textbf{...}
    parts = re.split(r'(\\textbf\{[^}]*\}|\\textit\{[^}]*\}|\\emph\{[^}]*\})', text)

    for part in parts:
        m_bold = re.match(r'\\textbf\{([^}]*)\}', part)
        m_italic = re.match(r'\\textit\{([^}]*)\}', part)
        m_emph = re.match(r'\\emph\{([^}]*)\}', part)

        if m_bold:
            run = p.add_run(clean_latex(m_bold.group(1)))
            run.bold = True
        elif m_italic or m_emph:
            content = (m_italic or m_emph).group(1)
            run = p.add_run(clean_latex(content))
            run.italic = True
        else:
            run = p.add_run(clean_latex(part))

        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        rpr = run._element.get_or_add_rPr()
        ea_font = rpr.makeelement(qn('w:rFonts'), {})
        ea_font.set(qn('w:eastAsia'), 'SimSun')
        rpr.insert(0, ea_font)

    return p


def parse_table(lines):
    """Parse a LaTeX tabular into a list of rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('\\') and not '&' in line:
            continue
        if '&' in line:
            # Remove trailing \\
            line = re.sub(r'\\\\.*$', '', line)
            cells = [clean_latex(c.strip()) for c in line.split('&')]
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows, caption=''):
    """Add a formatted table to the document."""
    if not rows:
        return

    if caption:
        add_formatted_paragraph(doc, caption, bold=True, font_size=10,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:
                    run.bold = True

    doc.add_paragraph()  # spacing


def build_docx():
    """Main conversion function."""
    with open(TEX_PATH, 'r', encoding='utf-8') as f:
        tex = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ================================================================
    # TITLE
    # ================================================================
    add_formatted_paragraph(
        doc,
        'A Transferable Deep Reinforcement Learning Framework for '
        'Farmland Spatial Layout Optimization Using Parcel-Level Scoring Policy',
        bold=True, font_size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12
    )

    # AUTHORS (anonymous for double-blind review)
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_auth.add_run('Anonymous authors')
    run_auth.font.name = 'Times New Roman'
    run_auth.font.size = Pt(12)
    p_auth.paragraph_format.space_after = Pt(12)

    # Target journal
    add_formatted_paragraph(
        doc, 'Target Journal: International Journal of Geographical Information Science',
        font_size=10, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )

    # ================================================================
    # ABSTRACT
    # ================================================================
    add_formatted_paragraph(doc, 'Abstract', bold=True, font_size=13, space_before=12, space_after=6)

    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', tex, re.DOTALL)
    if abstract_match:
        abstract_text = abstract_match.group(1).strip()
        add_formatted_paragraph(doc, abstract_text, font_size=10, space_after=6)

    # Keywords
    kw_text = 'deep reinforcement learning; parcel-level spatial optimization; farmland layout optimization; action masking; cross-dataset transfer; spatial planning'
    p = doc.add_paragraph()
    run_label = p.add_run('Keywords: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Times New Roman'
    run_kw = p.add_run(kw_text)
    run_kw.font.size = Pt(10)
    run_kw.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(18)

    # ================================================================
    # SECTION 1: Introduction
    # ================================================================
    add_formatted_paragraph(doc, '1. Introduction', bold=True, font_size=14, space_before=18, space_after=8)

    intro_text = [
        r"The spatial optimization of discrete land-use parcels---determining which spatial units to reclassify and in what sequence to improve aggregate landscape metrics---is a fundamental problem in geographical information science. It arises in diverse contexts including nature reserve site selection, afforestation planning, urban land-use allocation, and forest harvesting scheduling. A particularly pressing instance is the optimization of farmland spatial layouts in mountainous and hilly regions, where farmland fragmentation, steep slopes, and ecological constraints present compounding challenges \citep{liu2014land, long2018land}. In China, the policy of ``returning steep-slope cropland to forests'' (Grain for Green) encourages the exchange of high-slope farmland for low-slope forest land, aiming to consolidate arable resources on gentler terrain while restoring ecological function on steep slopes \citep{chen2020cropland}. Formulating an optimal parcel exchange plan requires balancing multiple competing spatial objectives---average slope reduction, spatial contiguity of arable blocks, and strict conservation of total farmland count---under combinatorial constraints whose search space grows super-exponentially with the number of parcels. This combination of discrete spatial units, neighborhood-dependent objectives, and global conservation constraints makes the problem representative of a broad class of parcel-level spatial layout optimization challenges in GIScience.",

        r"Traditional approaches to spatial land-use allocation have been studied extensively over the past two decades. Multi-objective evolutionary algorithms, particularly Genetic Algorithms (GA) and the Non-dominated Sorting Genetic Algorithm II (NSGA-II), have been widely adopted for generating Pareto-optimal land-use configurations \citep{cao2012spatial, stewart2004genetic, deb2002fast}. Integer and linear programming methods can guarantee global optimality for linearizable objectives but struggle with nonlinear spatial interactions \citep{aerts2003spatial}. Ant Colony Optimization (ACO) has been applied to land-use zoning by leveraging pheromone-based spatial search \citep{li2011coupling}, and its multi-type extension (MACO) was shown to handle large-area allocation problems effectively \citep{liu2012multi}. Simulated Annealing (SA) offers a single-solution metaheuristic alternative with theoretical convergence guarantees \citep{kirkpatrick1983optimization}. Cellular automata (CA) coupled with machine learning have also been used for simulating land-use change dynamics \citep{li2011coupling}. However, these methods share fundamental limitations: (1) the search space grows exponentially with the number of parcels, making fine-grained optimization at the parcel level (e.g., >10,000 parcels) computationally intractable for population-based methods; (2) they lack the ability to learn \emph{transferable spatial reasoning}---each new study area requires a completely independent optimization run, with no mechanism to carry learned spatial patterns from one geographic context to another; and (3) the quality of solutions depends heavily on hand-designed operators (crossover, mutation, pheromone update rules) that encode domain-specific heuristics rather than learning them from data.",

        r"Deep reinforcement learning (DRL) has emerged as a powerful paradigm for sequential decision-making and combinatorial optimization. Pioneering work by \citet{bello2016neural} demonstrated that policy gradient methods can learn competitive heuristics for the Traveling Salesman Problem (TSP) without supervised labels. Subsequent advances have extended DRL to vehicle routing \citep{nazari2018reinforcement}, job-shop scheduling \citep{zhang2020learning}, bin packing \citep{zhao2022learning}, and chip placement \citep{mirhoseini2021graph}. A common theme in these successes is that DRL can discover problem-specific strategies through interaction with a simulated environment, avoiding the need for explicit algorithmic design. In the domain of urban and spatial planning, DRL has seen emerging exploratory applications: \citet{wang2021deep} used adversarial learning for urban configuration generation, \citet{zheng2023spatial} applied DRL to spatial planning-aware landscape generation and urban community layouts. \citet{mao2019learning} demonstrated DRL for resource scheduling in computing clusters, illustrating the generality of DRL for allocation problems. However, the application of DRL to fine-grained parcel-level spatial layout optimization---where individual spatial units must be reassigned subject to conservation constraints and neighborhood-dependent objectives---remains largely unexplored. More critically, no existing work has addressed the challenge of \emph{cross-dataset transferability} in spatial optimization: training a DRL agent on one geographic area and deploying it on a structurally different area without retraining, a capability that would be transformative for GIS practitioners who routinely face the same class of optimization problem across many administrative units.",

        "This paper addresses these gaps by proposing a transferable DRL framework for parcel-level spatial layout optimization, demonstrated on the farmland--forest exchange problem. Our approach shifts the paradigm from purely numerical land-use quotas to spatially-explicit micro-level allocation, addressing a core challenge in modern GIScience. Our contributions are as follows:",
    ]

    for para in intro_text:
        add_rich_paragraph(doc, para)

    contributions = [
        "We formulate parcel-level spatial layout optimization as a Markov Decision Process (MDP) and solve it using Maskable PPO, a policy gradient method that respects invalid-action constraints through action masking. The MDP formulation naturally handles the sequential nature of parcel exchanges and the constraint that each spatial unit may be modified at most once.",
        "We design the ParcelScoringPolicy, a dimension-invariant neural architecture where a shared Multi-Layer Perceptron (MLP) independently scores each spatial unit by concatenating its local features with global state statistics. Because the network parameters are independent of the number of parcels N, the same trained weights can be directly applied to datasets of arbitrary size---enabling zero-shot cross-dataset transfer across geographically distinct areas.",
        "We propose a composite reward function for multi-objective spatial optimization under conservation constraints, with four synergistic components---slope reduction, contiguity improvement, quadratic count penalty, and pair-completion bonus---and demonstrate through comprehensive ablation that all components are essential for successful training. We further introduce a paired inference protocol that guarantees strict conservation at evaluation time.",
        "We conduct extensive experiments on real-world cadastral data (10,653 parcels for training, 8,185 for transfer validation) with comparisons against six baselines (Random, Greedy, GA, ACO, SA, NSGA-II), six-configuration ablation, four-parameter sensitivity analysis, and cross-dataset transfer experiments, providing the most comprehensive evaluation of DRL for parcel-level spatial optimization to date.",
    ]
    for i, c in enumerate(contributions, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(c)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # ================================================================
    # SECTION 2: Study Area and Data
    # ================================================================
    add_formatted_paragraph(doc, '2. Study Area and Data', bold=True, font_size=14, space_before=18, space_after=8)

    add_formatted_paragraph(doc, '2.1 Study area', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'Two villages in the hilly region of Sichuan--Chongqing, southwest China, are selected as study sites (Fig. 1). This region lies at the transition zone between the Sichuan Basin and the surrounding mountainous terrain, characterized by a subtropical monsoon climate with an annual mean temperature of 16--18\textdegree C and annual precipitation of 1,000--1,200 mm. The terrain features undulating low hills and shallow valleys with elevations ranging from 200 to 800 m a.s.l., creating a landscape where farmland, forest, and settlement parcels interleave at fine spatial scales.')

    add_rich_paragraph(doc, r'\textbf{Banzhu Village} (29.60\textdegree N, 106.14\textdegree E) serves as the primary training and evaluation site. The village covers approximately 710 ha and contains 10,653 land-use parcels, of which 6,737 (63.2%) are classified as farmland (including both paddy fields and dry cropland) and 2,515 (23.6%) as forest. The average farmland parcel size is approximately 667 m\u00b2. Farmland parcels exhibit a mean slope of 11.41\textdegree \pm 7.09\textdegree, with nearly half (49.2%) falling in the 6--15\textdegree range and 26.3% exceeding 15\textdegree---well above the threshold typically recommended for mechanized cultivation. Forest parcels are concentrated on steeper terrain with a mean slope of 16.24\textdegree \pm 7.84\textdegree. This inversion---where farmland occupies slopes steeper than necessary while some forest sits on relatively gentle terrain---creates a clear opportunity for land-type exchange.')

    add_rich_paragraph(doc, r'\textbf{Heping Village}, located in the same geographic region, serves as the cross-dataset transfer validation site. It covers approximately 546 ha with 8,185 parcels. Critically, Heping Village has a very different land-type composition: only 1,896 (23.2%) parcels are farmland while 4,354 (53.2%) are forest---roughly the inverse of Banzhu Village. The farmland mean slope is 10.14\textdegree \pm 9.41\textdegree and the forest mean slope is 19.00\textdegree \pm 11.20\textdegree. This contrasting composition makes Heping Village a stringent test for transfer learning, as the agent must generalize across a substantially different farmland-to-forest ratio.')

    # Figure 1
    fig1_path = os.path.join(SCRIPT_DIR, 'results_v7', 'figures', 'study_area_map.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_paragraph(doc,
        'Fig. 1. Location of the two study areas: (a) position within China; '
        '(b) Banzhu Village, the training site (10,653 parcels); '
        '(c) Heping Village, the transfer validation site (8,185 parcels). '
        'Parcels are colored by land-use type.',
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_formatted_paragraph(doc, '2.2 Data description', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'Parcel-level land-use data are stored as ESRI Shapefiles in the CGCS2000 3-Degree Gauss--Kruger Zone 35 projection (EPSG:4523). Each parcel record contains a grid cell identifier, land-use classification (DLMC), maximum slope value, perimeter, and area. Character encoding is UTF-8.')

    add_rich_paragraph(doc, r'The input datasets were constructed by fusing three publicly available data sources in ArcGIS Pro. Land-use classification was derived from the Sentinel-2 10 m Land Use/Land Cover Time Series (Esri, available via ArcGIS Living Atlas of the World), which provides globally consistent land-cover maps at 10 m resolution circa 2017. Terrain slope was computed from the Copernicus DEM GLO-30 (European Space Agency, accessed through Google Earth Engine), a global digital elevation model at 30 m posting. Administrative boundaries at the township level were obtained from the 1:5,000,000 Township Vector Boundaries of China (2020) published by the National Earth System Science Data Center. The three layers were intersected, clipped to the village boundaries, and rasterized into uniform grid cells of 25.82 m \u00d7 25.82 m (~667 m\u00b2), yielding the parcel-level polygons used throughout this study.')

    add_rich_paragraph(doc, 'Table 1 summarizes the two datasets.')

    # Table 1: Data statistics
    table1_rows = [
        ['Statistic', 'Banzhu Village', 'Heping Village'],
        ['Total parcels', '10,653', '8,185'],
        ['Farmland parcels', '6,737 (63.2%)', '1,896 (23.2%)'],
        ['Forest parcels', '2,515 (23.6%)', '4,354 (53.2%)'],
        ['Other parcels', '1,401 (13.2%)', '1,935 (23.6%)'],
        ['Swappable parcels', '9,252', '6,250'],
        ['Total area (ha)', '710.20', '545.67'],
        ['Mean farmland slope (\u00b0)', '11.41 \u00b1 7.09', '10.14 \u00b1 9.41'],
        ['Median farmland slope (\u00b0)', '10.34', '7.72'],
        ['Mean forest slope (\u00b0)', '16.24 \u00b1 7.84', '19.00 \u00b1 11.20'],
        ['Farmland contiguity index', '6.28', '5.24'],
        ['Farmland slope \u2264 6\u00b0', '1,652 (24.5%)', '831 (43.8%)'],
        ['Farmland slope 6\u201315\u00b0', '3,316 (49.2%)', '532 (28.1%)'],
        ['Farmland slope 15\u201325\u00b0', '1,447 (21.5%)', '363 (19.1%)'],
        ['Farmland slope > 25\u00b0', '322 (4.8%)', '170 (9.0%)'],
    ]
    add_table_to_doc(doc, table1_rows, 'Table 1. Descriptive statistics of the two study-area datasets.')

    # ================================================================
    # SECTION 3: Methodology
    # ================================================================
    add_formatted_paragraph(doc, '3. Methodology', bold=True, font_size=14, space_before=18, space_after=8)

    # 3.1 Problem definition (NEW)
    add_formatted_paragraph(doc, '3.1 Problem definition', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'Consider a study area partitioned into M land-use parcels P = {p_1, p_2, ..., p_M}. Each parcel p_i is characterized by a slope value s_i >= 0, an area a_i > 0, a set of spatial neighbors N(i), and a land-use type l_i in {F, L, O} (farmland, forest, or other). Only farmland and forest parcels are candidates for exchange; we denote the set of swappable parcels as P_s = {p_i | l_i in {F, L}} with N = |P_s|. Let P_F and P_L denote the initial farmland and forest sets, respectively.')

    add_rich_paragraph(doc, r"An exchange plan Pi = {(p_{a1}, p_{b1}), ..., (p_{aQ}, p_{bQ})} consists of Q farmland--forest pairs, where each pair specifies that parcel a_k is converted from farmland to forest and parcel b_k from forest to farmland. Each parcel may participate in at most one pair. After executing Pi, the new farmland set is P_F' = (P_F \ {p_{ak}}) union {p_{bk}}.")

    add_rich_paragraph(doc, r"The farmland spatial layout optimization problem seeks an exchange plan Pi* that solves the following bi-objective optimization:")

    add_formatted_paragraph(doc,
        'Pi* = argmin S(P_F\'),   argmax C(P_F\')',
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

    add_rich_paragraph(doc, 'subject to:')

    constraints = [
        "|P_F'| = |P_F|  (farmland count conservation)",
        "|Pi| <= Q_max  (exchange budget)",
        "p_{ak} != p_{aj}, p_{bk} != p_{bj} for all k != j  (uniqueness)",
    ]
    for c in constraints:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    add_rich_paragraph(doc, r"where S(P_F') = (1/|P_F'|) * sum(s_i) is the mean farmland slope, and C(P_F') = (1/|P_F'|) * sum(|N(i) intersect P_F'|) is the farmland contiguity index (mean number of farmland neighbors per farmland parcel). The exchange budget is set to Q_max = 100 (200 individual conversions), an empirical setting representing approximately 1.5% of the farmland parcel stock; this parameter is configurable and can be adjusted to match local policy quotas or planning scenarios.")

    add_rich_paragraph(doc, r"Note that because all parcels in our grid-based dataset have identical area (a_i ~ 667 m\u00b2, \u03c3 = 0), the farmland count conservation constraint is equivalent to farmland area conservation. In practice, if parcels have variable sizes, the count constraint should be replaced or supplemented by an area-based constraint.")

    add_rich_paragraph(doc, r'This problem is computationally hard: the number of feasible exchange plans is C(|P_F|, Q) * C(|P_L|, Q) * Q!, which for our training dataset (|P_F| = 6,737, |P_L| = 2,515, Q = 100) exceeds 10^500---far beyond the reach of exhaustive search. Moreover, the contiguity objective C() creates spatial dependencies between exchange decisions: the value of converting a particular parcel depends on which other parcels have already been converted, precluding simple greedy decomposition.')

    # 3.2 MDP formulation (was 3.1)
    add_formatted_paragraph(doc, '3.2 MDP formulation', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'To address the sequential and interdependent nature of the exchange decisions, we reformulate the optimization problem as a finite-horizon MDP (S, A, T, R, \gamma). The key insight is that each exchange pair can be decomposed into two sequential actions (one farmland-to-forest conversion and one forest-to-farmland conversion), naturally mapping to the agent--environment interaction loop of reinforcement learning.')

    add_rich_paragraph(doc, r'\textbf{State space.} At step t, the state s_t concatenates per-parcel features and global features. Each swappable parcel has K=6 features (normalized slope, current land type, neighborhood farmland ratio, neighborhood mean slope, normalized area, slope deviation from global mean), plus G=8 global features (normalized mean slope, contiguity, farmland count deviation, step progress, farmland ratio, forest ratio, slope change rate, contiguity change rate).')

    add_rich_paragraph(doc, r'\textbf{Action space.} Each action a_t selects a swappable parcel whose land type is flipped (farmland \to forest or forest \to farmland). An action mask ensures that each parcel is selected at most once per episode, and that the action is valid given the current state.')

    add_rich_paragraph(doc, r'\textbf{Reward function.} We design a composite reward with four terms:')

    # Reward equation as text
    add_formatted_paragraph(doc,
        'r_t = \u03bb\u2081 \u00b7 (s\u0304_{t-1} - s\u0304_t) / (|s\u0304_0| + \u03b5) '
        '+ \u03bb\u2082 \u00b7 (c_t - c_{t-1}) / (|c_0| + \u03b5) '
        '- \u03bb\u2083 \u00b7 (|n_t^f - n_0^f| / n_0^f)\u00b2 '
        '+ \u03b2 \u00b7 1[n_t^f = n_0^f]',
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

    reward_items = [
        'Slope reduction (\u03bb\u2081 = 1000): incentivizes converting high-slope farmland to forest.',
        'Contiguity improvement (\u03bb\u2082 = 500): rewards configurations that cluster farmland parcels.',
        'Quadratic count penalty (\u03bb\u2083 = 500): penalizes deviations from the initial farmland count, enforcing a soft conservation constraint.',
        'Pair-completion bonus (\u03b2 = 1.0): rewards the agent whenever the farmland count returns to its initial value, encouraging balanced pairwise exchanges.',
    ]
    for item in reward_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    add_rich_paragraph(doc, r'The denominators |s_0| + \epsilon and |c_0| + \epsilon in the reward function normalize each improvement term into percentage-of-initial-value form, ensuring that slope and contiguity contributions are expressed on comparable scales regardless of their original physical units (degrees vs. neighbor count). This normalization enables the hyperparameters \lambda_1 and \lambda_2 to be tuned on the same order of magnitude.')

    add_rich_paragraph(doc, r'\textbf{Episode structure.} Each episode comprises a fixed 200 steps with discount factor \gamma = 0.995. No early termination is used.')

    # 3.2 ParcelScoringPolicy
    add_formatted_paragraph(doc, '3.3 ParcelScoringPolicy', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'A key challenge is that different study areas have different numbers of swappable parcels N, making standard fixed-dimension policies non-transferable. We address this with the \textit{ParcelScoringPolicy} (Fig. 2), whose core idea is to evaluate each parcel independently with a shared network.')

    # Figure 2
    fig2_path = os.path.join(SCRIPT_DIR, 'results_v7', 'figures', 'architecture_diagram.png')
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_paragraph(doc,
        'Fig. 2. Architecture of the proposed Maskable PPO framework. The ParcelScoringPolicy '
        'scores each parcel independently using shared weights, making the network dimension-invariant '
        'with respect to the number of parcels N.',
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_rich_paragraph(doc, r'\textbf{Scorer network.} For each parcel i, the scorer takes as input the concatenation of its K-dimensional local features and the G-dimensional global features: logit_i = Scorer([f^(i); g]). The scorer is a two-layer MLP: (K+G=14) \to 128 \to 64 \to 1 with Tanh activations. Crucially, all N parcels share the same scorer weights. The logits are then passed through a masked softmax to obtain action probabilities: invalid actions (already-flipped parcels) are set to -\infty before the softmax.')

    add_rich_paragraph(doc, r'\textbf{Value network.} The value function depends only on global features: V(s_t) = Value(g_t), implemented as (G=8) \to 128 \to 64 \to 1 with Tanh activations.')

    add_rich_paragraph(doc, r'\textbf{Dimension invariance.} Because the scorer evaluates each parcel independently and the value network uses only global features, the total number of network parameters is independent of N. Trained weights from Banzhu Village (N = 9,252) can be loaded directly into a model for Heping Village (N = 6,250) with no architectural change and no fine-tuning.')

    add_rich_paragraph(doc, r'\textbf{Rationale for MLP over graph neural networks.} Because the land-parcel adjacency structure naturally defines a graph, one might consider graph neural networks (GCN, GAT) as an alternative to the MLP-based scorer. We deliberately chose the MLP design for three reasons. First, with N > 9,000 nodes and 200 RL steps per episode, recomputing graph convolutions at every step would incur substantial computational overhead; message-passing GNNs have O(N + |E|) per-layer cost, which must be repeated at each RL timestep as node features change after every exchange. Second, transferring a GNN across datasets requires not only weight sharing but also handling different graph topologies---variable-degree nodes, different connectivity patterns---adding implementation complexity. Third, our hand-crafted neighborhood features (neighborhood farmland ratio, neighborhood mean slope) already capture the essential one-hop spatial context that a GCN layer would learn, while the two-layer MLP provides nonlinear combinations of these features. The empirical results---0.41 s inference time, successful zero-shot transfer, and statistically significant contiguity improvement---suggest that the MLP with explicit neighborhood features is a pragmatic and effective design for this problem scale.')

    # 3.3 Training
    add_formatted_paragraph(doc, '3.4 Training with Maskable PPO', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'We use Maskable PPO from the sb3-contrib library [sb3contrib], which extends PPO [schulman2017proximal] to handle invalid-action masking. Key hyperparameters are listed in Table 2.')

    # Table 2: Hyperparameters
    table2_rows = [
        ['Hyperparameter', 'Value'],
        ['Algorithm', 'Maskable PPO'],
        ['Learning rate', '3 \u00d7 10\u207b\u2074 (Adam)'],
        ['Rollout length (n_steps)', '2,048'],
        ['Minibatch size', '256'],
        ['PPO epochs', '3'],
        ['Discount factor (\u03b3)', '0.995'],
        ['GAE parameter (\u03bb_GAE)', '0.95'],
        ['Clip range', '0.2'],
        ['Entropy coefficient', '0.01'],
        ['Total timesteps', '1,000,000'],
    ]
    add_table_to_doc(doc, table2_rows, 'Table 2. Training hyperparameters.')

    # 3.4 Paired inference
    add_formatted_paragraph(doc, '3.5 Paired inference', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'A fundamental challenge in applying DRL to constrained optimization is ensuring that hard constraints are satisfied during deployment. In our setting, the trained agent optimizes a composite reward that includes a soft farmland count penalty (Eq. 2), which encourages but does not guarantee farmland conservation. During unconstrained inference, the trained agent may convert more farmland to forest than vice versa, resulting in a net farmland loss. In our experiments, the unconstrained agent achieves a stronger slope reduction (-4.90%) but incurs a net loss of 130 farmland parcels (\u0394n^f = -130), which violates the strict conservation requirement of real-world planning.')

    add_rich_paragraph(doc, r'To address this, we introduce \textit{paired inference}, a post-training evaluation protocol that guarantees strict farmland conservation while fully leveraging the learned scoring function. The procedure operates as follows:')

    paired_steps = [
        'Farmland-to-forest step: The action mask is restricted to farmland-only parcels (parcels currently classified as farmland that have not yet been selected). The agent\'s scorer ranks all valid farmland parcels and selects the highest-scoring one for conversion to forest. This step reduces the farmland count by one.',
        'Forest-to-farmland step: The action mask is then restricted to forest-only parcels. The scorer ranks all valid forest parcels and selects the highest-scoring one for conversion to farmland. This step restores the farmland count.',
    ]
    for step in paired_steps:
        p = doc.add_paragraph(style='List Number')
        parts = step.split(': ', 1)
        run1 = p.add_run(parts[0] + ': ')
        run1.bold = True
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run2 = p.add_run(parts[1])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_rich_paragraph(doc, r'These two steps form one exchange pair. By strictly alternating between the two masks, the farmland count remains exactly at its initial value after every even-numbered step. We execute 100 such pairs (200 individual conversions), matching the episode length used during training. This design decouples the learning of parcel quality assessment (during training, where soft penalties allow exploration) from the enforcement of hard constraints (during inference, where paired masks guarantee conservation).')

    add_rich_paragraph(doc, r'\textbf{Discussion of the soft-to-hard constraint gap.} We acknowledge that paired inference is an engineering protocol rather than a theoretically principled solution to the constraint satisfaction problem. The gap between the soft quadratic penalty used during training and the hard conservation requirement of deployment reflects a fundamental tension in applying standard DRL to constrained optimization. In principle, constrained MDPs with Lagrangian dual variables could learn to satisfy the conservation constraint directly during training, eliminating the need for post-hoc enforcement. We adopt the paired inference approach for two pragmatic reasons: (1) Maskable PPO with soft penalties is straightforward to implement using existing libraries (sb3-contrib), whereas constrained policy optimization requires custom algorithmic development; and (2) the scorer trained with soft penalties still learns a high-quality parcel ranking function, and the paired protocol merely imposes a deterministic scheduling layer on top of this learned ranking. The empirical success---zero net farmland loss with competitive optimization performance---validates this pragmatic decomposition, but a principled constrained MDP formulation remains an important direction for future work (Section 6.7).')

    # ================================================================
    # SECTION 4: Experiments
    # ================================================================
    add_formatted_paragraph(doc, '4. Experiments', bold=True, font_size=14, space_before=18, space_after=8)

    add_formatted_paragraph(doc, '4.1 Experimental setup', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, r'All experiments were conducted on a consumer-grade PC (Intel Core, 32 GB RAM, no dedicated GPU; PyTorch on CPU). Training one model for 1,000,000 timesteps takes approximately 4 hours. Evaluation metrics include: (1) mean farmland slope change and its percentage; (2) contiguity index change; (3) farmland count change; and (4) computational time. All stochastic methods are run with 5 random seeds and we report mean \pm standard deviation.')

    add_formatted_paragraph(doc, '4.2 Baseline methods', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'We compare against six baselines spanning the full spectrum of optimization paradigms:')

    baselines = [
        'Random: Randomly selects 100 farmland\u2013forest pairs for exchange.',
        'Greedy: Each step selects the pair that maximizes \u03bb\u2081 \u00b7 \u0394slope + \u03bb\u2082 \u00b7 \u0394contiguity.',
        'GA: Genetic Algorithm with population 100, 500 generations, tournament selection, set-based crossover, and elitism [cao2012spatial].',
        'ACO: Ant Colony Optimization with 30 ants, 200 iterations, pheromone-guided parcel selection, and candidate list acceleration [li2011coupling, liu2012multi].',
        'SA: Simulated Annealing with exponential cooling and periodic reheating.',
        'NSGA-II: Non-dominated Sorting Genetic Algorithm II with two objectives (slope improvement, contiguity improvement), population 100, 500 generations.',
    ]
    for b in baselines:
        p = doc.add_paragraph(style='List Bullet')
        # Bold the method name
        parts = b.split(': ', 1)
        run1 = p.add_run(parts[0] + ': ')
        run1.bold = True
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run2 = p.add_run(parts[1])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    add_rich_paragraph(doc, 'All baselines use the same constraint: exactly 100 paired exchanges (200 individual conversions), conserving total farmland count.')

    add_rich_paragraph(doc, 'To ensure fair comparison, all population-based baselines (GA, ACO, SA, NSGA-II) employ a customized pairwise mutation operator that randomly selects one farmland parcel and one forest parcel and swaps their labels, guaranteeing exact farmland count conservation throughout the search. This design aligns the action space of all methods with the DRL agent\'s pairwise exchange mechanism, so that observed performance differences reflect the quality of parcel selection rather than differences in constraint handling.')

    # 4.3 Ablation
    add_formatted_paragraph(doc, '4.3 Ablation study', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'We train six model variants to assess the contribution of each component (Table 3):')

    table3_rows = [
        ['Configuration', 'Modification', 'Validation target'],
        ['Full Model', 'All default settings', 'Baseline reference'],
        ['w/o Pair Bonus', '\u03b2 = 0', 'Pair mechanism'],
        ['w/o Contiguity', '\u03bb\u2082 = 0', 'Contiguity objective'],
        ['w/o Count Penalty', '\u03bb\u2083 = 0', 'Conservation constraint'],
        ['Small Network', 'Scorer/Value: [64, 32]', 'Network capacity'],
        ['Short Horizon', '\u03b3 = 0.99', 'Temporal discounting'],
    ]
    add_table_to_doc(doc, table3_rows, 'Table 3. Ablation study configurations.')

    # 4.4 Sensitivity
    add_formatted_paragraph(doc, '4.4 Hyperparameter sensitivity', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'We vary one parameter at a time while holding others at their defaults: \u03bb\u2081 \u2208 {500, 1000, 2000}, \u03bb\u2083 \u2208 {200, 500, 1000}, \u03b2 \u2208 {0.5, 1.0, 2.0}, and \u03b3 \u2208 {0.99, 0.995, 0.999}, yielding 12 configurations (7 newly trained + 5 reused from ablation and default).')

    # 4.5 Transfer
    add_formatted_paragraph(doc, '4.5 Cross-dataset transfer', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'To evaluate transferability, we load the scorer and value network weights trained on Banzhu Village and apply them directly to Heping Village without any retraining or fine-tuning. We compare against three models trained from scratch on Heping Village (seeds 0, 1, 2) using identical hyperparameters.')

    # ================================================================
    # SECTION 5: Results
    # ================================================================
    add_formatted_paragraph(doc, '5. Results', bold=True, font_size=14, space_before=18, space_after=8)

    # 5.1 Training convergence (NEW)
    add_formatted_paragraph(doc, '5.1 Training convergence', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'Fig. 3 shows the training dynamics over 1,000,000 timesteps (~5,000 episodes). The episodic reward exhibits three distinct phases: (1) an initial exploration phase (0--200,000 steps) where the agent discovers the exchange mechanism, with rewards climbing from approximately -25 to +10; (2) a rapid improvement phase (200,000--600,000 steps) where the agent refines its parcel selection strategy, reaching rewards of +35; and (3) a convergence plateau (600,000--1,000,000 steps) where rewards stabilize around +38--+39. The mean farmland slope decreases from 11.41\textdegree to approximately 10.86\textdegree in the training environment, indicating that the agent learns to consistently identify beneficial exchanges. The farmland contiguity index improves from 6.28 to approximately 6.28, reflecting the agent\'s attention to spatial clustering.')

    # Figure 3: Training curves
    fig_train_path = os.path.join(SCRIPT_DIR, 'results_v7', 'training_curves.png')
    if os.path.exists(fig_train_path):
        doc.add_picture(fig_train_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_paragraph(doc,
        'Fig. 3. Training curves over 1,000,000 timesteps: (a) episodic reward; '
        '(b) mean farmland slope; (c) farmland contiguity index. '
        'Shaded areas represent rolling standard deviation.',
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 5.2 Multi-seed DRL performance (NEW)
    add_formatted_paragraph(doc, '5.2 Multi-seed DRL performance', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'To assess the robustness of the DRL agent, we train five models with different random seeds and evaluate each using paired inference. Table 4 reports per-seed results. Four out of five seeds achieve highly consistent performance (slope ~ -3.7%, contiguity ~ +0.05), while seed 1 shows a somewhat weaker slope reduction (-2.20%) but the strongest contiguity improvement (+0.065), suggesting a different exploration trajectory. The mean slope of parcels converted out of farmland (to forest) is consistently high (33.6--33.7\textdegree), while the mean slope of parcels converted into farmland is low (4.8--16.8\textdegree), confirming that the agent learns the intended ``high-slope-out, low-slope-in\'\' strategy.')

    # Table 4: Multi-seed results
    table_multiseed_rows = [
        ['Seed', 'Slope \u0394 (%)', 'Contiguity \u0394', 'FC', 'Slope Out (\u00b0)', 'Slope In (\u00b0)', 'Time (s)'],
        ['0', '-3.69', '+0.054', '0', '33.72', '5.35', '0.71'],
        ['1', '-2.20', '+0.065', '0', '33.71', '16.79', '0.58'],
        ['2', '-3.72', '+0.051', '0', '33.56', '4.96', '0.65'],
        ['3', '-3.76', '+0.048', '0', '33.71', '4.83', '0.66'],
        ['4', '-3.69', '+0.054', '0', '33.72', '5.37', '0.69'],
        ['Mean \u00b1 std', '-3.41 \u00b1 0.60', '+0.054 \u00b1 0.006', '0', '33.68 \u00b1 0.07', '7.46 \u00b1 5.05', '0.66 \u00b1 0.05'],
    ]
    add_table_to_doc(doc, table_multiseed_rows, 'Table 4. Per-seed DRL results with paired inference on Banzhu Village.')

    # 5.3 Baseline comparison
    add_formatted_paragraph(doc, '5.3 Baseline comparison', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'Table 5 presents the comprehensive comparison of all methods on Banzhu Village.')

    # Table 5: Comparison
    table5_rows = [
        ['Method', 'Slope Change (%)', 'Contiguity Change', 'FC', 'Time (s)'],
        ['Random', '+0.61 \u00b1 0.11', '-0.113 \u00b1 0.007', '0', '<0.01'],
        ['GA', '-2.28 \u00b1 0.07', '-0.017 \u00b1 0.007', '0', '26.4 \u00b1 0.7'],
        ['ACO', '-2.48 \u00b1 0.03', '+0.004 \u00b1 0.003', '0', '3.2 \u00b1 0.6'],
        ['SA', '-3.28 \u00b1 0.08', '+0.017 \u00b1 0.005', '0', '79.6 \u00b1 3.7'],
        ['Greedy\u2020', '-3.78', '+0.044', '0', '<0.01'],
        ['NSGA-II', '-1.84 \u00b1 0.07', '+0.011 \u00b1 0.006', '0', '134.7 \u00b1 19.0'],
        ['DRL-paired (ours)', '-3.41 \u00b1 0.60', '+0.054 \u00b1 0.006', '0', '0.41'],
    ]
    add_table_to_doc(doc, table5_rows, 'Table 5. Performance comparison of all optimization methods on Banzhu Village. DRL results are from 5-seed paired inference. \u2020: single deterministic run.')

    add_rich_paragraph(doc, r'Among all methods, Greedy achieves the best single-metric slope reduction (-3.78%), followed by our DRL approach (-3.41%) and SA (-3.28%). However, DRL achieves the best contiguity improvement (+0.054) among all methods---23% better than Greedy (+0.044), the next-best method. GA (-2.28%), ACO (-2.48%), and NSGA-II (-1.84%) achieve moderate slope reductions but with weaker or negligible contiguity gains. The Random baseline worsens both objectives (slope +0.61%, contiguity -0.113), confirming the non-trivial nature of the optimization task.')

    add_rich_paragraph(doc, r"The key distinction of DRL lies in its simultaneous optimization of both objectives. Greedy's superior slope performance results from its exhaustive O(N\u00b2) search that selects the locally optimal exchange at each step; however, this single-step greedy strategy sacrifices long-term spatial structure, yielding lower contiguity improvement. SA achieves comparable slope reduction (-3.28%) but requires nearly 80 s of computation per run. NSGA-II, despite being explicitly designed for multi-objective optimization, achieves the weakest slope reduction (-1.84%) among the non-random baselines, likely because its population is distributed across the Pareto front rather than concentrated in the high-performing region. ACO and GA fall in between, with ACO showing a slight edge in contiguity (+0.004 vs. -0.017) due to its pheromone-guided spatial search.")

    add_rich_paragraph(doc, r"No single method dominates all others across every metric. Greedy leads on slope, DRL leads on contiguity, and SA offers a balanced intermediate profile. The choice of method depends on the planner's priorities: if slope reduction alone is the primary objective, Greedy provides a strong deterministic solution; if multi-objective balance and transferability are valued, DRL offers the most attractive trade-off.")

    add_rich_paragraph(doc, r'\textbf{Statistical significance.} We conduct pairwise Mann--Whitney U tests (Table 6). The contiguity improvement of DRL-paired is statistically significant over all baselines (p < 0.01) with large effect sizes (Cohen\'s d = 6.5--24.7). For slope, DRL is significantly better than Random (p < 0.01) but not statistically distinguishable from SA, GA, or ACO due to the variance across seeds and the relatively small sample size (n = 5).')

    # Table 6: Significance
    table6_rows = [
        ['Comparison', 'Metric', 'p-value', 'Significant?', "Cohen's d"],
        ['vs. Random', 'Slope', '0.008', 'Yes', '8.26'],
        ['', 'Contiguity', '0.008', 'Yes', '24.67'],
        ['vs. GA', 'Slope', '0.095', 'No', '2.34'],
        ['', 'Contiguity', '0.008', 'Yes', '10.42'],
        ['vs. ACO', 'Slope', '0.151', 'No', '1.94'],
        ['', 'Contiguity', '0.008', 'Yes', '9.81'],
        ['vs. SA', 'Slope', '0.151', 'No', '0.29'],
        ['', 'Contiguity', '0.008', 'Yes', '6.49'],
    ]
    add_table_to_doc(doc, table6_rows, 'Table 6. Mann\u2013Whitney U test results: DRL-paired vs. baselines.')

    add_rich_paragraph(doc, r'\textbf{Computational efficiency.} After a one-time training cost of approximately 4 hours, the DRL agent optimizes a new layout in 0.41 s---8\u00d7 faster than ACO, 64\u00d7 faster than GA, 191\u00d7 faster than SA, and 328\u00d7 faster than NSGA-II. This amortized inference speed is critical for interactive planning scenarios where planners may need to evaluate many candidate layouts.')

    # 5.4 Ablation
    add_formatted_paragraph(doc, '5.4 Ablation study', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, r'Table 7 presents the ablation results. Removing any of the three reward components (pair bonus, contiguity reward, or count penalty) leads to complete model failure: slope increases by 0.16--0.19% and contiguity decreases by 0.08--0.09. This demonstrates that all three components are essential and synergistic.')

    table7_rows = [
        ['Configuration', 'Slope Change (%)', 'Contiguity Change', 'FC'],
        ['Full Model', '-3.69', '+0.054', '0'],
        ['w/o Pair Bonus (\u03b2=0)', '+0.19', '-0.082', '0'],
        ['w/o Contiguity (\u03bb\u2082=0)', '+0.16', '-0.090', '0'],
        ['w/o Count Penalty (\u03bb\u2083=0)', '+0.19', '-0.082', '0'],
        ['Small Network [64, 32]', '-3.55', '+0.061', '0'],
        ['Short Horizon (\u03b3=0.99)', '-2.40', '+0.062', '0'],
    ]
    add_table_to_doc(doc, table7_rows, 'Table 7. Ablation study results. Models evaluated with paired inference on Banzhu Village.')

    add_rich_paragraph(doc, r'Architectural changes are less sensitive: the small network ([64, 32] instead of [128, 64]) retains 96% of the slope optimization (-3.55% vs. -3.69%), indicating that the model is not capacity-limited. Reducing the discount factor from 0.995 to 0.99 (Short Horizon) causes a moderate degradation to -2.40% slope reduction (-35% relative drop), suggesting that long-horizon planning is important for optimizing the 200-step exchange sequence.')

    # 5.5 Sensitivity
    add_formatted_paragraph(doc, '5.5 Hyperparameter sensitivity', bold=True, font_size=12, space_before=12, space_after=6)

    # Figure 4: Sensitivity
    fig_sens_path = os.path.join(SCRIPT_DIR, 'results_v7', 'sensitivity', 'sensitivity_analysis.png')
    if os.path.exists(fig_sens_path):
        add_rich_paragraph(doc, 'Fig. 4 summarizes the sensitivity analysis results.')
        doc.add_picture(fig_sens_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc,
            'Fig. 4. Hyperparameter sensitivity analysis. Each panel varies one parameter '
            'while holding others at their defaults (marked with blue squares). '
            'Top row: slope change (%); bottom row: contiguity change.',
            font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_rich_paragraph(doc, 'Several patterns emerge:')

    sensitivity_items = [
        'Slope weight (\u03bb\u2081) exhibits a narrow optimal region around the default value of 1,000. Both \u03bb\u2081 = 500 (slope +0.72%) and \u03bb\u2081 = 2,000 (slope +0.16%) cause model failure. A lower weight provides insufficient gradient signal for slope optimization, while an excessively high weight destabilizes training by creating reward magnitudes that dominate the other objectives.',
        'Count penalty (\u03bb\u2083) is more robust: \u03bb\u2083 = 1,000 produces slightly better slope reduction (-3.80%) than the default (-3.69%), though contiguity drops from +0.054 to +0.044. A penalty of \u03bb\u2083 = 200 is insufficient.',
        'Pair bonus (\u03b2) shows a threshold between 0.5 (failure) and 1.0 (success). The value \u03b2 = 2.0 performs comparably, with slightly improved contiguity (+0.061 vs. +0.054) at a small cost in slope (-3.59% vs. -3.69%).',
        'Discount factor (\u03b3) is sensitive at both ends: \u03b3 = 0.99 causes a moderate drop (-2.40%), while \u03b3 = 0.999 causes complete failure (+0.19%). The value \u03b3 = 0.995 provides the best balance between short-term exploitation and long-horizon planning.',
    ]
    for i, item in enumerate(sensitivity_items, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # 5.6 Transfer
    add_formatted_paragraph(doc, '5.6 Cross-dataset transfer', bold=True, font_size=12, space_before=12, space_after=6)
    add_rich_paragraph(doc, 'Table 8 presents the cross-dataset results on Heping Village.')

    table8_rows = [
        ['Method', 'Slope Change (%)', 'Contiguity Change', 'Training Time'],
        ['Transfer (Banzhu \u2192 Heping)', '-15.99', '+0.108', '0 h'],
        ['From Scratch (seed 0)', '-0.00', '-0.341', '2.7 h'],
        ['From Scratch (seed 1)', '-0.00', '-0.423', '2.7 h'],
        ['From Scratch (seed 2)', '-0.01', '-0.312', '2.7 h'],
        ['From Scratch (mean \u00b1 std)', '-0.00 \u00b1 0.01', '-0.359 \u00b1 0.047', '2.7 h'],
    ]
    add_table_to_doc(doc, table8_rows, 'Table 8. Cross-dataset comparison on Heping Village: transfer learning vs. training from scratch.')

    add_rich_paragraph(doc, r'The transferred model achieves a striking -15.99% farmland slope reduction and +0.108 contiguity improvement---far exceeding its performance on the training site (-3.69%). This is attributable to Heping Village\'s much higher forest-to-farmland ratio (4,354 vs. 1,896), providing a richer pool of low-slope forest parcels available for exchange. In contrast, all three from-scratch models achieve near-zero slope improvement and severely degrade contiguity (-0.36 \pm 0.05), indicating that the reward landscape of Heping Village is more challenging to learn from scratch with the default hyperparameters.')

    add_rich_paragraph(doc, 'This result provides strong evidence for the practical value of transfer learning: a model trained on one representative village can immediately optimize layouts for other sites in the same geographic region, eliminating the need for site-specific training.')

    # ================================================================
    # SECTION 6: Discussion
    # ================================================================
    add_formatted_paragraph(doc, '6. Discussion', bold=True, font_size=14, space_before=18, space_after=8)

    add_formatted_paragraph(doc, '6.1 Advantages of DRL over traditional methods', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'Our results demonstrate several qualitative advantages of DRL for land-use optimization that distinguish it from both heuristic and metaheuristic approaches.')

    add_rich_paragraph(doc, r'First, the learned policy captures implicit spatial reasoning. Rather than relying on hand-crafted heuristics (Greedy) or population-based random search (GA, ACO), the DRL agent learns to consider neighborhood context and global state when selecting parcels. The six per-parcel features---particularly the neighborhood farmland ratio and neighborhood mean slope---allow the scorer to assess each parcel in its spatial context. This is reflected in the superior contiguity improvement (+0.054), which is statistically significant over all baselines (p < 0.01), including methods that explicitly incorporate spatial information in their fitness functions.')

    add_rich_paragraph(doc, r'Second, the once-trained model provides amortized instant inference. While the initial training investment is substantial (~6.5 h on CPU), inference completes in 0.41 s per layout---two orders of magnitude faster than SA (79.6 s) and three orders faster than NSGA-II (134.7 s). In practical planning scenarios where decision-makers need to evaluate dozens or hundreds of candidate configurations under different policy constraints, this speed advantage becomes decisive. The total compute cost of DRL becomes favorable over GA after fewer than 100 inference queries.')

    add_rich_paragraph(doc, r'Third, the dimension-invariant architecture enables zero-shot transfer---a capability fundamentally absent from all baseline methods. GA, ACO, SA, and NSGA-II must each run a complete optimization from scratch for every new study area. In contrast, the DRL agent trained on Banzhu Village can be deployed on Heping Village (with a different number of parcels and a very different farmland-to-forest ratio) by simply loading the trained weights into a new environment wrapper.')

    # 6.2 Comparison with existing DRL approaches (NEW)
    add_formatted_paragraph(doc, '6.2 Comparison with existing DRL approaches in spatial planning', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'While DRL has been applied to urban planning tasks [wang2021deep, zheng2023spatial], our work differs in several important aspects. First, urban planning DRL typically operates on grid cells or districts with relatively low spatial resolution (hundreds to low thousands of units), whereas our framework handles 10,653 individual parcels---creating an action space an order of magnitude larger. The ParcelScoringPolicy addresses this scalability challenge through its dimension-invariant design, avoiding the quadratic parameter growth that would occur with a standard fully-connected policy.')

    add_rich_paragraph(doc, r'Second, existing spatial planning DRL work has not addressed the transferability challenge. [wang2021deep] train and evaluate on the same city grid, and [zheng2023spatial] use fixed-size landscape patches. Our framework is, to our knowledge, the first to demonstrate successful zero-shot transfer between geographically distinct areas with different spatial extents and land-type compositions.')

    add_rich_paragraph(doc, r'Third, the constrained optimization aspect---maintaining exact farmland count conservation---is unique to our problem domain. Urban planning DRL typically optimizes unconstrained objectives (e.g., accessibility, land-use diversity). Our paired inference protocol provides a principled solution for combining learned policies with hard constraints.')

    # 6.3 Role of reward engineering (expanded)
    add_formatted_paragraph(doc, '6.3 Role of reward engineering', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'The ablation study reveals that all four reward components interact synergistically, and the removal of any one of the three key components (pair bonus, contiguity reward, count penalty) causes complete model failure---not merely degraded performance but a reversal of optimization direction (slope increases by 0.16--0.19%). This fragility has important implications for reward engineering in DRL for constrained optimization.')

    add_rich_paragraph(doc, r'The pair-completion bonus (\u03b2) is particularly critical. Without it, the agent has no incentive to perform balanced farmland--forest exchanges; it may convert multiple farmland parcels in sequence without compensating forest-to-farmland conversions, leading to a policy that the quadratic penalty alone cannot correct during training. The sensitivity analysis confirms this: \u03b2 = 0.5 (insufficient bonus) produces the same failure mode as \u03b2 = 0 (ablation), while \u03b2 >= 1.0 enables successful learning. This suggests a phase-transition-like behavior where the pair bonus must exceed a threshold to guide exploration toward balanced exchange strategies.')

    add_rich_paragraph(doc, r'The contiguity reward (\u03bb\u2082) serves a dual role. Beyond its direct contribution to contiguity optimization, it provides a spatially informative gradient signal that helps the agent learn to consider parcel neighborhoods. Without contiguity reward, the agent\'s scoring function degenerates into a pure slope ranker that ignores spatial context.')

    add_rich_paragraph(doc, r'The quadratic (rather than linear) form of the count penalty is essential for maintaining a smooth gradient near the conservation target n_t^f = n_0^f. A linear penalty would create a constant gradient regardless of the deviation magnitude, while the quadratic form penalizes large deviations increasingly harshly while providing a gentle gradient near the target. These findings underscore the importance of careful reward engineering in DRL for constrained optimization and suggest that similar multi-component reward structures may be necessary for other spatial optimization problems.')

    add_rich_paragraph(doc, r'\textbf{Multi-objective scalability: a cross-paradigm challenge.} It is important to recognize that the reward engineering fragility observed here is not a limitation unique to DRL, but rather a manifestation of the intrinsic difficulty of multi-objective optimization that presents itself differently across paradigms. Multi-objective evolutionary algorithms such as NSGA-II elegantly sidestep the weight-setting problem by maintaining a Pareto front of non-dominated solutions; however, as the number of objectives grows, the Pareto front becomes a high-dimensional surface, requiring exponentially larger populations to approximate adequately, and the selection pressure toward any particular region of the front weakens\u2014a phenomenon known as the "curse of dimensionality in objective space." In our experiments, NSGA-II already achieves the weakest slope reduction (-1.84%) among non-random baselines with only two objectives. ACO faces an analogous challenge: in the single-objective setting, pheromone accumulation on high-quality exchange pairs provides an effective search guidance mechanism. When extended to multiple objectives, ACO must either combine objectives into a single fitness for pheromone update\u2014reducing to the same weighted-sum problem as DRL\u2014or maintain separate pheromone matrices for each objective (the MACO approach), which introduces inter-matrix coordination complexity that grows with the number of objectives. Mathematical programming (LP/IP) handles constraints and multiple objectives through well-established duality theory, but requires that objectives and constraints be linearizable\u2014a condition violated by spatially dependent objectives like our contiguity index, which involves counting farmland neighbors of farmland parcels and thus exhibits quadratic combinatorial structure. SA collapses multiple objectives into a single acceptance criterion, encountering the same weight-sensitivity problem as DRL. In summary, every optimization paradigm encounters an equivalent form of multi-objective difficulty; the particular manifestation varies (weight sensitivity in DRL and SA, population scaling in NSGA-II, pheromone coordination in ACO, linearization requirements in mathematical programming), but the fundamental tension between conflicting objectives is inescapable. The advantage of the DRL framework lies not in avoiding this difficulty, but in offering unique compensating strengths\u2014scalability to >10,000 parcels, sub-second amortized inference, and zero-shot cross-dataset transfer\u2014that no competing paradigm provides. Future integration of constrained MDP formulations with Lagrangian dual variables for automatic weight adjustment represents a promising path toward combining DRL\'s scalability advantages with more principled multi-objective handling.')

    # 6.4 Transfer mechanism (expanded)
    add_formatted_paragraph(doc, '6.4 Transfer mechanism', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'The success of cross-dataset transfer arises from the design of the ParcelScoringPolicy. Because the scorer evaluates each parcel using only local features (K = 6) and global statistics (G = 8)---both of which are normalized to [0, 1] or centered around zero---the learned representation captures relative parcel quality rather than absolute coordinates or dataset-specific patterns. In essence, the scorer learns a general principle: ``prefer high-slope farmland for conversion to forest, and prefer low-slope forest near existing farmland for conversion.\'\' This principle generalizes across villages with different sizes, shapes, and farmland/forest distributions.')

    add_rich_paragraph(doc, r'The transfer results are particularly striking: the transferred model achieves a -15.99% slope reduction on Heping Village, compared to -3.69% on its training site (Banzhu Village). This amplified performance is explained by Heping Village\'s favorable land-type ratio: with 4,354 forest parcels but only 1,896 farmland parcels, the pool of candidate low-slope forest parcels is large relative to the number of exchanges needed (100 pairs), giving the scorer abundant high-quality options to choose from.')

    add_rich_paragraph(doc, r'The complete failure of from-scratch training on Heping Village (~0% slope improvement, contiguity degradation of -0.36 \pm 0.05) across all three seeds provides strong negative evidence that reinforces the value of transfer. We identify three compounding factors behind this failure. First, the skewed land-type distribution (farmland:forest ~ 1:2.3 in Heping vs. 2.7:1 in Banzhu) reduces the density of farmland parcels in the action space, making the reward signal from slope reduction sparser and noisier. Second, the 200-step episode length constitutes a much larger fraction of the farmland count in Heping (200/1,896 ~ 10.5%) than in Banzhu (200/6,737 ~ 3.0%), amplifying the destabilizing effect of farmland count deviations on the quadratic penalty term and creating a harsher optimization landscape. Third, the hyperparameters (\u03bb\u2081, \u03bb\u2082, \u03bb\u2083, \u03b2, \u03b3) were tuned on Banzhu Village; our sensitivity analysis shows that \u03bb\u2081 and \u03b3 have very narrow optimal ranges, and it is likely that these values are suboptimal for Heping\'s different reward dynamics. This suggests that from-scratch training on Heping would require a dedicated hyperparameter search---precisely the kind of per-site engineering effort that transfer learning eliminates. Transfer learning circumvents this cold-start problem entirely by providing a pre-trained scorer that immediately identifies beneficial exchanges based on general spatial principles learned from the training site.')

    # 6.5 Sensitivity (expanded)
    add_formatted_paragraph(doc, '6.5 Sensitivity of hyperparameters', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'The sensitivity analysis reveals that \u03bb\u2081 (slope weight) and \u03b3 (discount factor) require careful tuning, with narrow optimal ranges. Both halving and doubling \u03bb\u2081 from its default of 1,000 causes complete model failure---a sensitivity pattern more reminiscent of learning rate tuning in deep learning than of typical optimization parameter selection. This is because \u03bb\u2081 controls not only the objective weighting but also the magnitude of gradient signals during policy updates: too small and the slope objective is drowned out by the count penalty; too large and the slope term dominates, destabilizing the multi-objective balance.')

    add_rich_paragraph(doc, r'The discount factor \u03b3 shows a similarly narrow window: \u03b3 = 0.99 yields moderate performance (-2.40% slope) while \u03b3 = 0.999 causes complete failure. With \u03b3 = 0.999, the effective planning horizon extends well beyond the 200-step episode, causing the value function to overweight distant returns and leading to credit assignment difficulties. The optimal \u03b3 = 0.995 corresponds to an effective horizon of approximately 1/(1-\u03b3) = 200 steps, exactly matching the episode length.')

    add_rich_paragraph(doc, r'In contrast, \u03bb\u2083 (count penalty) and \u03b2 (pair bonus) are more robust above their respective thresholds. The count penalty \u03bb\u2083 = 1,000 even slightly outperforms the default on slope (-3.80% vs. -3.69%), suggesting that stronger conservation enforcement can be beneficial. In practice, we recommend starting from the default values reported here and exploring \u03bb\u2083 and \u03b2 first when adapting to new study areas, as these parameters are more forgiving of adjustment.')

    # 6.6 Fairness of baseline comparison (NEW)
    add_formatted_paragraph(doc, '6.6 Fairness of baseline comparison', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, 'A rigorous evaluation requires transparent discussion of comparison fairness. Several aspects warrant examination.')

    add_rich_paragraph(doc, r'\textbf{Exchange budget control.} All methods perform exactly 100 paired exchanges (200 individual swaps), ensuring that the quantity of land-type changes is held constant. This "equal exchange count" protocol means that every method operates under the same physical constraint---the same number of parcels are affected. It is therefore the quality of parcel selection that differentiates the methods, not the volume of modification.')

    add_rich_paragraph(doc, r'\textbf{Computational budget.} The methods vary substantially in wall-clock time: Greedy takes <0.01 s (deterministic), DRL inference takes 0.41 s, ACO takes 3.2 s, GA takes 26.4 s, SA takes 79.6 s, and NSGA-II takes 134.7 s. DRL\'s training phase (~4 h) is not included in inference time but must be acknowledged as a one-time upfront cost. We emphasize that comparisons are controlled by exchange count, not by compute time. If one were to grant all methods equal compute budgets (e.g., by running more GA/SA iterations), their results might improve. However, the exchange-count protocol better reflects the practical planning constraint: the number of land parcels that can be realistically reclassified is limited by administrative and ecological considerations, not by available CPU time.')

    add_rich_paragraph(doc, r'\textbf{Hyperparameter tuning.} The baseline methods (GA, ACO, SA, NSGA-II) use standard parameter settings from the literature (e.g., population size = 100, crossover rate = 0.8 for GA; 30 ants, \u03b1 = 1.0, \u03b2 = 2.0 for ACO). These parameters were not exhaustively tuned for our specific problem. It is possible that dedicated parameter sweeps could improve baseline performance. Conversely, the DRL approach also uses default hyperparameters from Stable-Baselines3 with minimal tuning, and the sensitivity analysis shows that several DRL hyperparameters have narrow optimal ranges, suggesting that our DRL results may also not represent the theoretical maximum.')

    add_rich_paragraph(doc, r"\textbf{Positioning DRL's contribution.} Given the above considerations, we do not claim that DRL achieves absolute superiority on all metrics. Indeed, Greedy produces better slope reduction (-3.78% vs. -3.41%), and SA achieves comparable slope performance. The core contribution of the DRL framework lies in three complementary advantages that no single baseline method offers simultaneously: (a) best multi-objective performance, with the highest contiguity improvement among all methods; (b) sub-second inference after training, enabling interactive planning workflows; and (c) zero-shot transferability to new study areas without re-optimization---a capability absent from all baseline approaches.")

    # 6.7 Independent validation via FFI
    add_formatted_paragraph(doc, '6.7 Independent validation via Farmland Fragmentation Index', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, r'To validate that DRL optimization produces genuine spatial layout improvements beyond the training objectives (slope and contiguity), we employ the Farmland Fragmentation Index (FFI) [niu2023ffi] as an independent post-hoc evaluation metric. FFI is a composite landscape index that integrates six indicators across three dimensions---Patch Scale (PS: number of patches NP, largest patch index LPI), Shape Regularity (SR: landscape shape index LSI, area-weighted mean shape index AWMSI), and Spatial Distribution (SD: patch density PD, aggregation index AI)---combined via AHP-derived weights (W_PS=0.21, W_SR=0.24, W_SD=0.55). FFI ranges from 0 to 1, with higher values indicating greater farmland fragmentation. Crucially, FFI is \emph{not} part of the reward function and was not used during training; any improvement reflects emergent spatial quality gains.')

    # FFI table
    ffi_header = ['Dataset', 'FFI (before)', 'FFI (after)', '\u0394FFI', '\u0394AI (%)']
    ffi_data = [
        ['Banzhu Village (5 seeds)', '0.6017', '0.6001 \u00b1 0.0002', '-0.0016 \u00b1 0.0002', '+0.34'],
        ['Heping Village (transfer)', '0.6260', '0.6234', '-0.0026', '+0.57'],
    ]
    add_table_to_doc(doc, [ffi_header] + ffi_data,
                     caption='Table 9. Farmland Fragmentation Index (FFI) before and after DRL optimization. Lower FFI indicates less fragmentation. AI denotes the aggregation index, the dominant driver of FFI change.')

    add_rich_paragraph(doc, r'On Banzhu Village, FFI decreases from 0.6017 to 0.6001 \u00b1 0.0002 across five seeds (\u0394FFI = -0.0016), driven primarily by an increase in the aggregation index (AI: +0.34%). On Heping Village, the cross-dataset transferred model achieves a larger FFI reduction (\u0394FFI = -0.0026, AI: +0.57%), consistent with its stronger slope and contiguity improvements.')

    add_rich_paragraph(doc, r'We note that the FFI changes, while consistently negative (indicating reduced fragmentation), are modest in absolute magnitude. This is expected for two reasons: (1) only 100 paired exchanges are performed out of thousands of parcels, limiting the achievable spatial reorganization; and (2) the grid-based parcels already exhibit high baseline aggregation (AI \u2248 40%), leaving limited room for further improvement. Nevertheless, the directional consistency across all seeds and both study sites---and the larger effect on the transfer site---provides independent evidence that the learned policy produces spatially coherent layout improvements rather than merely optimizing the reward signal.')

    # 6.8 Limitations and future work (expanded)
    add_formatted_paragraph(doc, '6.8 Limitations and future work', bold=True, font_size=12, space_before=12, space_after=6)

    add_rich_paragraph(doc, 'Several limitations merit discussion.')

    limitations = [
        'The current framework considers only binary land-type exchanges (farmland \u2194 forest). Extending to multi-class exchanges (e.g., including grassland, water bodies, or construction land) would broaden applicability but significantly increase the action space and complicate the conservation constraints, as multiple land-type budgets must be maintained simultaneously.',
        'Only two optimization objectives (slope and contiguity) are considered. Real-world farmland consolidation planning also involves accessibility (distance to roads and settlements), irrigation infrastructure proximity, soil quality indices, and ecological corridor connectivity. While adding new input features to the scorer is straightforward (simply expanding K from 6 to a larger dimension), adding new optimization objectives to the reward function poses a more fundamental challenge. Our sensitivity analysis already reveals that the slope weight \u03bb\u2081 has a very narrow optimal range; introducing additional reward terms (e.g., \u03bb\u2084\u00b7\u0394soil + \u03bb\u2085\u00b7\u0394accessibility) would expand the hyperparameter search space combinatorially and increase the risk of inter-objective gradient conflicts\u2014where improving one objective worsens another through competing gradient signals. The ablation study further underscores this fragility: removing any single reward component causes complete model failure, indicating that the current multi-term balance is already delicate. Scaling to five or more objectives under the weighted-sum scalarization paradigm would likely require prohibitive tuning effort.',
        'The from-scratch training failure on Heping Village suggests that the reward function and hyperparameters may require site-specific tuning when the land-type ratio deviates substantially from the training distribution, though transfer learning effectively mitigates this limitation.',
        'The study uses grid-based parcels of approximately uniform size (667 m\u00b2); validation on irregularly shaped real cadastral parcels with highly variable areas is needed.',
        'The current evaluation is limited to two villages in the same geographic region; cross-region transfer (e.g., from hilly terrain to plains) remains to be investigated.',
    ]
    for i, lim in enumerate(limitations, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'({i}) {lim}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_rich_paragraph(doc, 'Future directions include: (a) multi-agent DRL for coordinating exchanges across adjacent administrative units, where inter-village boundary effects create additional optimization opportunities; (b) incorporating richer geospatial features via graph neural networks that can explicitly model the parcel adjacency structure; (c) integrating the framework into interactive GIS tools for real-time planning support, leveraging the sub-second inference speed; (d) extending the transfer mechanism to include few-shot fine-tuning with limited target-site data; (e) applying the framework to other spatial resource allocation problems in natural resource management; and (f) addressing the multi-objective scalability limitation through alternative optimization paradigms. Promising approaches include constrained MDPs, where secondary objectives (e.g., soil quality, road accessibility) are enforced as hard constraints via extended action masks rather than as reward terms, thereby avoiding additional weight tuning; multi-objective RL (MORL), which maintains a set of policies along the Pareto front rather than collapsing objectives into a single scalar; and lexicographic optimization, which orders objectives by priority and optimizes each sequentially without degrading higher-priority outcomes. These approaches would enable the framework to accommodate the richer multi-criteria requirements of real-world land-use planning without the reward engineering fragility observed in the current weighted-sum formulation.')

    add_rich_paragraph(doc, r'\textbf{Broader applicability to natural resource management.} The core design of our framework---parcel-level scoring with dimension-invariant weights, multi-objective reward engineering, and paired inference for hard constraint satisfaction---is not specific to farmland layout optimization. Several natural resource management problems share the same structural characteristics (large-scale combinatorial parcel selection, spatial contiguity objectives, and conservation constraints) and could benefit from the proposed approach.')

    add_rich_paragraph(doc, r'Nature reserve site selection is a representative example. Selecting parcels for a protected area to maximize habitat connectivity under a budget constraint is structurally analogous to our problem: parcels must be chosen from a large candidate set, spatial contiguity (habitat corridor connectivity) is a key objective, and the total protected area is bounded. The ParcelScoringPolicy could be directly adapted by redefining local features (e.g., species richness, habitat quality) and global features (e.g., total connectivity index, budget utilization) while retaining the same network architecture.')

    add_rich_paragraph(doc, r'Afforestation and reforestation planning presents another natural extension. Selecting optimal planting sites to maximize carbon sequestration while satisfying soil, slope, and water-availability constraints mirrors the exchange-plan formulation: each candidate site can be scored by a shared MLP using local terrain and soil features concatenated with global carbon-budget statistics. The dimension-invariant property would enable models trained in one watershed to be transferred to adjacent watersheds without retraining.')

    add_rich_paragraph(doc, r'More broadly, any spatial resource allocation problem that involves (i) selecting or reassigning discrete spatial units from a large candidate set, (ii) optimizing objectives that depend on the spatial neighborhood structure, and (iii) satisfying global budget or conservation constraints fits naturally into the MDP formulation and ParcelScoringPolicy architecture proposed in this work. Examples include water resource allocation across irrigation districts, forest harvesting scheduling with adjacency buffer constraints, and renewable energy facility siting under ecological sensitivity constraints. The common thread is that the scorer learns a relative quality assessment of spatial units that generalizes across problem instances, enabling rapid inference and cross-site transfer---advantages that are particularly valuable in natural resource management, where planners routinely face the same class of optimization problem across many geographically distinct administrative units.')

    # ================================================================
    # SECTION 7: Conclusions
    # ================================================================
    add_formatted_paragraph(doc, '7. Conclusions', bold=True, font_size=14, space_before=18, space_after=8)

    add_rich_paragraph(doc, 'We presented a deep reinforcement learning framework for parcel-level spatial layout optimization that formulates farmland--forest exchange planning as an MDP solved with Maskable PPO. The core innovation---the ParcelScoringPolicy---enables dimension-invariant inference and zero-shot cross-dataset transfer across geographically distinct areas. Our main findings are:')

    conclusions = [
        'Effective multi-objective optimization. On Banzhu Village (10,653 parcels), the method achieves -3.41% \u00b1 0.60% mean farmland slope reduction and +0.054 \u00b1 0.006 contiguity improvement with zero farmland loss, outperforming all baselines in contiguity (p < 0.01).',
        'Critical reward components. Ablation confirms that the pair bonus, contiguity reward, and count penalty are all essential; removing any one causes complete model failure.',
        'Robust transferability. When applied to Heping Village (8,185 parcels) without retraining, the model achieves -15.99% slope reduction\u2014far exceeding models trained from scratch on the same site (\u22480%).',
        'Computational efficiency. After a one-time training investment of ~4 hours, inference takes 0.41 s, 8\u2013328\u00d7 faster than evolutionary baselines.',
        'Independent validation. The Farmland Fragmentation Index (FFI), a composite landscape metric not included in the reward function, consistently decreases after optimization (\u0394FFI = -0.0016 on Banzhu Village, -0.0026 on Heping Village), confirming that the learned policy produces genuine spatial layout improvements beyond the training objectives.',
    ]
    for i, c in enumerate(conclusions, 1):
        p = doc.add_paragraph(style='List Number')
        # Bold the title part
        parts = c.split('. ', 1)
        run1 = p.add_run(parts[0] + '. ')
        run1.bold = True
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(11)
        run2 = p.add_run(parts[1])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_rich_paragraph(doc, 'The framework offers a general-purpose tool for parcel-level spatial optimization that extends beyond the farmland application demonstrated here. Any spatial resource allocation problem involving discrete spatial units, neighborhood-dependent objectives, and global conservation constraints---such as nature reserve site selection, afforestation planning, or renewable energy facility siting---can be formulated within the same MDP and ParcelScoringPolicy architecture. The dimension-invariant design and zero-shot transferability are particularly valuable for GIS practitioners who routinely face the same class of optimization problem across many administrative units. In the agricultural domain specifically, by consolidating farmland onto gentler, more contiguous terrain, the optimized layouts directly support agricultural mechanization and scale farming, bridging the gap between computational optimization research and the practical demands of modern agricultural production.')

    # ================================================================
    # CRediT, Declaration, Data Availability
    # ================================================================
    # NOTE: CRediT and Acknowledgments are omitted for double-blind review.

    add_formatted_paragraph(doc, 'Declaration of Competing Interest', bold=True, font_size=14, space_before=18, space_after=8)
    add_rich_paragraph(doc, 'The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.')

    add_formatted_paragraph(doc, 'Data and Code Availability', bold=True, font_size=14, space_before=18, space_after=8)
    add_rich_paragraph(doc, 'The code, trained model weights, and evaluation scripts supporting this study have been anonymized and will be made publicly available upon acceptance. For peer review, an anonymized repository is provided at [Anonymous GitHub link to be inserted]. The cadastral datasets were derived from publicly available remote sensing products (Sentinel-2 10 m LULC via ArcGIS Living Atlas; Copernicus DEM GLO-30 via Google Earth Engine) and government-published administrative boundary data (National Earth System Science Data Center). The processed village-level shapefiles can be made available upon reasonable request to the corresponding author.')

    # ================================================================
    # REFERENCES
    # ================================================================
    add_formatted_paragraph(doc, 'References', bold=True, font_size=14, space_before=18, space_after=8)

    references = [
        'Aerts, J.C., Eisinger, E., Heuvelink, G.B., Stewart, T.J., 2003. Using linear integer programming for multi-site land-use allocation. Geographical Analysis 35(2), 148\u2013169.',
        'Bello, I., Pham, H., Le, Q.V., Norouzi, M., Bengio, S., 2016. Neural combinatorial optimization with reinforcement learning. arXiv preprint arXiv:1611.09940.',
        'Cao, K., Huang, B., Wang, S., Lin, H., 2012. Sustainable land use optimization using Boundary-based Fast Genetic Algorithm. Computers, Environment and Urban Systems 36(3), 257\u2013269.',
        'Chen, Y., Wang, Y., Wang, J., 2020. Cropland protection and rational use in China. Land Use Policy 99, 104905.',
        'Deb, K., Pratap, A., Agarwal, S., Meyarivan, T., 2002. A fast and elitist multiobjective genetic algorithm: NSGA-II. IEEE Transactions on Evolutionary Computation 6(2), 182\u2013197.',
        'Kirkpatrick, S., Gelatt, C.D., Vecchi, M.P., 1983. Optimization by simulated annealing. Science 220(4598), 671\u2013680.',
        'Li, X., Lao, C.H., Liu, X.P., Chen, Y.M., 2011. Coupling urban cellular automata with ant colony optimization for zoning protected natural areas under a changing landscape. International Journal of Geographical Information Science 25(4), 575\u2013593.',
        'Liu, X.P., Li, X., Shi, X., Huang, K., Liu, Y.L., 2012. A multi-type ant colony optimization (MACO) method for optimal land use allocation in large areas. International Journal of Geographical Information Science 26(7), 1325\u20131343.',
        'Liu, Y., Fang, F., Li, Y., 2014. Key issues of land use in China and implications for policy making. Land Use Policy 40, 6\u201312.',
        'Long, H., Qu, Y., Tu, S., Zhang, Y., Jiang, Y., 2018. Development of land use transitions research in China. Journal of Geographical Sciences 28(3), 375\u2013390.',
        'Mao, H., Schwarzkopf, M., Venkatakrishnan, S.B., Meng, Z., Alizadeh, M., 2019. Learning scheduling algorithms for data processing clusters. In: Proc. ACM SIGCOMM, pp. 270\u2013288.',
        'Mirhoseini, A., Goldie, A., Yazgan, M., et al., 2021. A graph placement methodology for fast chip design. Nature 594(7862), 207\u2013212.',
        'Nazari, M., Oroojlooy, A., Snyder, L., Takac, M., 2018. Reinforcement learning for solving the vehicle routing problem. In: NeurIPS, pp. 9839\u20139849.',
        'Niu, B., Ge, D., Yan, R., Ma, Y., Sun, D., Lu, Y., 2023. Investigation and comparison of spatial\u2013temporal characteristics of farmland fragmentation in China. Land 12(11), 2047.',
        'Schulman, J., Wolski, F., Dhariwal, P., Radford, A., Klimov, O., 2017. Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.',
        'Stable-Baselines3 Contributors, 2021. Stable-Baselines3 Contrib. https://github.com/Stable-Baselines-Team/stable-baselines3-contrib.',
        'Stewart, T.J., Janssen, R., van Herwijnen, M., 2004. A genetic algorithm approach to multiobjective land use planning. Computers & Operations Research 31(14), 2293\u20132313.',
        'Wang, D., Fu, Y., Wang, P., Huang, B., Lu, C.-T., 2021. Automated urban planning for reimagining city configuration via adversarial learning: quantification, generation, and evaluation. ACM Transactions on Spatial Algorithms and Systems 7(4), 1\u201328.',
        'Zhang, C., Song, W., Cao, Z., Zhang, J., Tan, P.S., Xu, C., 2020. Learning to dispatch for job shop scheduling via deep reinforcement learning. In: NeurIPS, pp. 1621\u20131632.',
        'Zhao, H., She, Q., Zhu, C., Yang, Y., Xu, K., 2022. Online 3D bin packing with constrained deep reinforcement learning. In: Proc. AAAI, pp. 741\u2013749.',
        'Zheng, Y., Lin, Y., Zhao, L., Wu, T., Jin, D., Li, Y., 2023. Spatial planning-aware landscape generation. In: Proc. AAAI, pp. 15500\u201315508.',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] {ref}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)

    # ================================================================
    # TITLE PAGE (for double-blind submission: placed at the end)
    # ================================================================
    # Page break
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    p_break = doc.add_paragraph()
    run_break = p_break.add_run()
    br = OxmlElement('w:br')
    br.set(_qn('w:type'), 'page')
    run_break._element.append(br)

    add_formatted_paragraph(
        doc, 'TITLE PAGE',
        bold=True, font_size=16,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18
    )

    add_formatted_paragraph(
        doc,
        'A Transferable Deep Reinforcement Learning Framework for '
        'Farmland Spatial Layout Optimization Using Parcel-Level Scoring Policy',
        bold=True, font_size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18
    )

    # Full author information
    p_tp_auth = doc.add_paragraph()
    p_tp_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tp_auth1 = p_tp_auth.add_run('Ning Zhou')
    run_tp_auth1.font.name = 'Times New Roman'
    run_tp_auth1.font.size = Pt(12)
    run_tp_sup1 = p_tp_auth.add_run('a')
    run_tp_sup1.font.name = 'Times New Roman'
    run_tp_sup1.font.size = Pt(12)
    run_tp_sup1.font.superscript = True
    run_tp_and = p_tp_auth.add_run(' and Xiang Jing')
    run_tp_and.font.name = 'Times New Roman'
    run_tp_and.font.size = Pt(12)
    run_tp_sup2 = p_tp_auth.add_run('a')
    run_tp_sup2.font.name = 'Times New Roman'
    run_tp_sup2.font.size = Pt(12)
    run_tp_sup2.font.superscript = True
    run_tp_star = p_tp_auth.add_run('*')
    run_tp_star.font.name = 'Times New Roman'
    run_tp_star.font.size = Pt(12)
    run_tp_star.font.superscript = True
    p_tp_auth.paragraph_format.space_after = Pt(12)

    # Affiliation
    p_tp_aff = doc.add_paragraph()
    p_tp_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tp_aff_sup = p_tp_aff.add_run('a')
    run_tp_aff_sup.font.name = 'Times New Roman'
    run_tp_aff_sup.font.size = Pt(11)
    run_tp_aff_sup.font.superscript = True
    run_tp_aff_text = p_tp_aff.add_run(' School of Software and Microelectronics, Peking University, Beijing 100871, China')
    run_tp_aff_text.font.name = 'Times New Roman'
    run_tp_aff_text.font.size = Pt(11)
    run_tp_aff_text.italic = True
    p_tp_aff.paragraph_format.space_after = Pt(12)

    # Corresponding author
    add_formatted_paragraph(
        doc, '*Corresponding author. Email: jingxiang@pku.edu.cn',
        font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )

    # CRediT Authorship Contribution Statement
    add_formatted_paragraph(doc, 'CRediT Authorship Contribution Statement', bold=True, font_size=13, space_before=12, space_after=8)

    p_tp_credit = doc.add_paragraph()
    run_tp_nz_name = p_tp_credit.add_run('Ning Zhou: ')
    run_tp_nz_name.bold = True
    run_tp_nz_name.font.name = 'Times New Roman'
    run_tp_nz_name.font.size = Pt(11)
    run_tp_nz_roles = p_tp_credit.add_run('Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Writing -- Original Draft, Visualization. ')
    run_tp_nz_roles.font.name = 'Times New Roman'
    run_tp_nz_roles.font.size = Pt(11)
    run_tp_xj_name = p_tp_credit.add_run('Xiang Jing: ')
    run_tp_xj_name.bold = True
    run_tp_xj_name.font.name = 'Times New Roman'
    run_tp_xj_name.font.size = Pt(11)
    run_tp_xj_roles = p_tp_credit.add_run('Supervision, Writing -- Review & Editing, Funding acquisition.')
    run_tp_xj_roles.font.name = 'Times New Roman'
    run_tp_xj_roles.font.size = Pt(11)
    p_tp_credit.paragraph_format.space_after = Pt(6)

    # Save
    doc.save(OUT_PATH)
    print(f"Word document saved: {OUT_PATH}")
    print(f"File size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")


if __name__ == '__main__':
    build_docx()
