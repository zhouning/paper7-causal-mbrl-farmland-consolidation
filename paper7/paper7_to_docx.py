# -*- coding: utf-8 -*-
"""Convert paper7_dreamer_farmland.tex to Word (.docx) format — English version."""

import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tex_to_docx import clean_latex

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEX_PATH = os.path.join(SCRIPT_DIR, 'paper7_dreamer_farmland.tex')
OUT_EN = os.path.join(SCRIPT_DIR, 'paper7_dreamer_farmland.docx')


def process_table(table_tex, doc):
    cap_m = re.search(r'\\caption\{(.+?)\}(?:\s*\\label)?', table_tex, re.DOTALL)
    caption = clean_latex(cap_m.group(1)) if cap_m else 'Table'

    tab_m = re.search(r'\\begin\{tabular\}\{([^}]+)\}(.+?)\\end\{tabular\}', table_tex, re.DOTALL)
    if not tab_m:
        doc.add_paragraph(f'[Table: {caption}]')
        return

    content = tab_m.group(2)
    rows = []
    for line in content.split('\\\\'):
        line = line.strip()
        line = re.sub(r'\\(toprule|midrule|bottomrule|hline|cline\{[^}]+\})', '', line).strip()
        if not line:
            continue
        line = re.sub(r'\\multirow\{[^}]+\}\{[^}]+\}', '', line)
        cells = [clean_latex(c.strip()) for c in line.split('&')]
        if any(c.strip() for c in cells):
            rows.append(cells)

    if not rows:
        doc.add_paragraph(f'[Table: {caption}]')
        return

    p = doc.add_paragraph()
    run = p.add_run(f'Table: {caption}')
    run.bold = True

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                table.rows[i].cells[j].text = cell_text

    doc.add_paragraph('')


def process_figure(fig_tex, doc):
    cap_m = re.search(r'\\caption\{(.+?)\}(?:\s*\\label)?', fig_tex, re.DOTALL)
    caption = clean_latex(cap_m.group(1)) if cap_m else 'Figure'

    img_m = re.search(r'\\includegraphics(?:\[.*?\])?\{(.+?)\}', fig_tex)
    img_path = img_m.group(1) if img_m else None

    if img_path:
        full_path = os.path.join(SCRIPT_DIR, img_path.replace('/', os.sep))
        if os.path.exists(full_path):
            doc.add_picture(full_path, width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f'[Image: {img_path}]')
    else:
        doc.add_paragraph('[Figure]')

    p = doc.add_paragraph()
    run = p.add_run(f'Figure: {caption}')
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')


def process_block(block_text, doc):
    block_text = block_text.strip()
    if not block_text:
        return

    env_pattern = re.compile(r'\\begin\{(table|figure)\}\[?\w*\]?.*?\\end\{\1\}', re.DOTALL)
    env_matches = list(env_pattern.finditer(block_text))

    if not env_matches:
        process_text(block_text, doc)
        return

    last_end = 0
    for m in env_matches:
        before = block_text[last_end:m.start()].strip()
        if before:
            process_text(before, doc)
        if m.group(1) == 'table':
            process_table(m.group(0), doc)
        else:
            process_figure(m.group(0), doc)
        last_end = m.end()

    after = block_text[last_end:].strip()
    if after:
        process_text(after, doc)


def process_text(text, doc):
    text = re.sub(r'\\label\{[^}]+\}', '', text)

    list_pattern = re.compile(r'\\begin\{(enumerate|itemize)\}(.+?)\\end\{\1\}', re.DOTALL)
    parts = list_pattern.split(text)

    idx = 0
    while idx < len(parts):
        part = parts[idx].strip()
        if part in ('enumerate', 'itemize'):
            list_content = parts[idx + 1] if idx + 1 < len(parts) else ''
            items = re.split(r'\\item(?:\[.*?\])?\s+', list_content)
            item_num = 0
            for item in items:
                item = item.strip()
                if not item:
                    continue
                item_num += 1
                style = 'List Number' if part == 'enumerate' else 'List Bullet'
                doc.add_paragraph(clean_latex(item), style=style)
            idx += 2
        else:
            part = re.sub(r'\\(begin|end)\{[^}]+\}', '', part).strip()
            if part:
                paras = re.split(r'\n\s*\n', part)
                for para in paras:
                    para = para.strip()
                    if not para:
                        continue
                    para = re.sub(r'\\begin\{algorithm\}.*?\\end\{algorithm\}',
                                  '[Algorithm - see LaTeX source]', para, flags=re.DOTALL)
                    para = re.sub(r'\\begin\{(equation|align)\*?\}(.+?)\\end\{\1\*?\}',
                                  lambda m: '[Eq] ' + clean_latex(m.group(2).strip()),
                                  para, flags=re.DOTALL)
                    cleaned = clean_latex(para)
                    cleaned = ' '.join(cleaned.split())
                    if cleaned:
                        doc.add_paragraph(cleaned)
            idx += 1


def build_docx(tex_content, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title_m = re.search(r'\\title\{(.+?)\}', tex_content, re.DOTALL)
    title_text = clean_latex(title_m.group(1)) if title_m else 'Paper 7'
    p = doc.add_heading(title_text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    doc.add_paragraph('Ning Zhou, Xiang Jing*').alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('School of Software and Microelectronics, Peking University, Beijing 100871, China')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('* Corresponding author: jingxiang@pku.edu.cn')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    abs_m = re.search(r'\\begin\{abstract\}(.+?)\\end\{abstract\}', tex_content, re.DOTALL)
    if abs_m:
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(clean_latex(abs_m.group(1).strip()))

    # Body
    body_start_m = re.search(r'\\section\{', tex_content)
    body_end_m = re.search(r'\\begin\{thebibliography\}', tex_content)
    if body_start_m and body_end_m:
        body = tex_content[body_start_m.start():body_end_m.start()]
    elif body_start_m:
        body = tex_content[body_start_m.start():]
    else:
        body = ''

    section_re = re.compile(r'\\(sub(?:sub)?)?section\*?\{([^}]+)\}')
    headers = [(m.start(), m.end(), m.group(1), m.group(2)) for m in section_re.finditer(body)]

    sec_num = 0
    subsec_num = 0
    for i, (start, end, stype, title) in enumerate(headers):
        if stype is None:
            sec_num += 1
            subsec_num = 0
            doc.add_heading(f'{sec_num}. {clean_latex(title)}', level=1)
        elif stype == 'sub':
            subsec_num += 1
            doc.add_heading(f'{sec_num}.{subsec_num} {clean_latex(title)}', level=2)

        next_start = headers[i + 1][0] if i + 1 < len(headers) else len(body)
        content = body[end:next_start]
        process_block(content, doc)

    # References
    doc.add_heading('References', level=1)
    bib_m = re.search(r'\\begin\{thebibliography\}.+?\\end\{thebibliography\}', tex_content, re.DOTALL)
    if bib_m:
        bib_text = bib_m.group(0)
        items = re.findall(
            r'\\bibitem\[([^\]]+)\]\{[^}]+\}\s*(.+?)(?=\\bibitem|\\end\{thebibliography\})',
            bib_text, re.DOTALL
        )
        for label, content in items:
            content = clean_latex(content.strip())
            content = re.sub(r'\\newblock\s*', '', content)
            content = ' '.join(content.split())
            doc.add_paragraph(f'[{clean_latex(label)}] {content}')

    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    with open(TEX_PATH, 'r', encoding='utf-8') as f:
        tex = f.read()

    build_docx(tex, OUT_EN)
    print('Done!')
