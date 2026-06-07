# -*- coding: utf-8 -*-
"""Generate fully Chinese version of Paper 7 as Word (.docx) format."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(SCRIPT_DIR, 'paper7_dreamer_farmland_cn.docx')


def set_cn_font(run, size=11, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    rpr = run._element.get_or_add_rPr()
    ea = rpr.makeelement(qn('w:rFonts'), {})
    ea.set(qn('w:eastAsia'), 'SimSun')
    rpr.insert(0, ea)


def add_para(doc, text, size=11, bold=False, italic=False,
             align=None, space_after=6, space_before=0, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size, bold, italic)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p


def add_heading_cn(doc, text, level=1, space_before=18, space_after=8):
    sizes = {1: 14, 2: 12, 3: 11}
    add_para(doc, text, size=sizes.get(level, 11), bold=True,
             space_before=space_before, space_after=space_after)


def add_table(doc, caption, headers, rows):
    add_para(doc, caption, size=10, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=4)
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_cn_font(run, size=9, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i + 1, j)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                set_cn_font(run, size=9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_list(doc, items, numbered=True):
    for i, item in enumerate(items, 1):
        prefix = f'{i}. ' if numbered else '\u2022 '
        p = doc.add_paragraph()
        run = p.add_run(prefix + item)
        set_cn_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.8)


# === Content ===

def write_frontmatter(doc):
    add_para(doc,
             '因果校准的基于模型强化学习耕地整治优化方法：基于学习环境动力学',
             size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, '周宁\u1d43，景翔\u1d43*',
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, '*通讯作者。电子邮件: jingxiang@pku.edu.cn',
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, '\u1d43 北京大学软件与微电子学院，北京 100871',
             size=10, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)


def write_abstract(doc):
    add_heading_cn(doc, '摘要')
    add_para(doc,
        '面向空间规划的深度强化学习（DRL）需要昂贵的环境交互——'
        '县域级耕地整治优化在A100 GPU上每次训练需8-12小时，'
        '且40%的训练因策略向奖励模型利用漂移而失败。'
        '本文提出一种基于模型的方法，用轻量级神经转移模型（237K参数）替代真实环境，'
        '使策略训练完全在CPU上45分钟内完成。'
        '转移模型从多样化行为策略收集的12,000条轨迹样本中学习状态动力学，'
        '与真实状态转移的余弦相似度达0.9998。'
        '一种新颖的因果奖励校准机制解决了奖励利用问题：'
        '利用真实轨迹数据的倾向得分匹配平均处理效应（ATT）估计'
        '修正学习奖励函数中的系统性偏差，'
        '产生校准因子以缩放高质量与低质量动作之间的预测奖励差异。'
        '在13个乡镇（2,600个区块、52,515个地块）的县域级耕地整治实验中，'
        '使用15个随机种子验证了三项发现：'
        '(1) 基于模型的策略在真实环境上实现-0.976%\u00b10.129%的坡度降低，'
        '超过在A100上训练12小时的无模型MARL（-0.84%\u00b10.07%）；'
        '(2) 因果校准进一步将性能提升至-1.102%\u00b10.100%，'
        '统计显著提升13.0%（p=0.004，Mann-Whitney U检验）；'
        '(3) 因果推断得到的校准因子（\u03b1=0.185）与网格搜索最优值（\u03b1=0.200）'
        '仅差3%，证明因果推断提供了免调参的校准替代方案。'
        '完整流程无需GPU，在普通硬件上2小时内完成。',
        size=10, space_after=6)
    p = doc.add_paragraph()
    run = p.add_run('关键词: ')
    set_cn_font(run, size=10, bold=True)
    run2 = p.add_run('基于模型的强化学习; 因果奖励校准; 耕地整治; '
                      '转移模型; 倾向得分匹配; 空间规划')
    set_cn_font(run2, size=10)
    p.paragraph_format.space_after = Pt(18)


def write_body(doc):
    write_sec1(doc)
    write_sec2(doc)
    write_sec3(doc)
    write_sec4(doc)
    write_sec5(doc)
    write_sec6(doc)
    write_sec7(doc)


def write_sec1(doc):
    add_heading_cn(doc, '1. 引言')
    add_para(doc,
        '面向空间规划问题的强化学习面临根本性的计算瓶颈：环境模拟代价高昂。'
        '在耕地整治优化中，每个环境步骤需要更新地块级土地利用状态、'
        '重新计算空间邻接指标并评估百亩方等策略目标。'
        '训练单个策略需500,000步，在A100 GPU上消耗8-12小时，且25-40%因策略漂移失败。')
    add_para(doc,
        '基于模型的强化学习（MBRL）提供原则性解决方案：'
        '学习近似环境状态动力学的转移模型，在学习模型上训练策略。'
        '然而MBRL尚未应用于大规模空间规划问题。')
    add_para(doc, '本文做出三项贡献：')
    add_list(doc, [
        '面向空间规划的学习环境：237K参数转移模型，余弦相似度0.9998，CPU 60分钟训练。',
        '因果奖励校准：修正5.4倍奖励高估，提升策略质量13.0%（p=0.004），'
        '因果推断\u03b1与网格搜索最优仅差3%。',
        '数量级计算缩减：完整流程CPU 2小时，无需GPU，15种子统计显著。',
    ])


def write_sec2(doc):
    add_heading_cn(doc, '2. 相关工作')
    add_heading_cn(doc, '2.1 基于模型的强化学习', level=2)
    add_para(doc,
        'MBRL方法通过学习环境动力学来降低样本复杂度。'
        'World Models引入了在学习的隐动力学模型生成的"梦境"轨迹上训练策略的概念。'
        'Dreamer V3将此扩展至多样化的连续控制领域。'
        'MuZero通过学习模型进行规划，在棋类和Atari中实现了超人性能。'
        '在空间领域，Zheng等将无模型DRL应用于城市社区规划，'
        '但尚无先前工作将MBRL应用于大规模空间规划。')
    add_para(doc,
        '本文方法与Dreamer风格方法的关键区别在于：'
        '我们不联合学习隐状态表示和动力学模型，'
        '而是直接在环境的原始观测空间中操作。'
        '这是可行的，因为耕地优化环境已提供了紧凑、信息丰富的状态表示'
        '（每区块17个特征\u00d72,600个区块+12个全局特征），'
        '无需学习编码。')
    add_heading_cn(doc, '2.2 奖励模型利用', level=2)
    add_para(doc,
        '智能体利用学习奖励模型的问题已有充分记录。'
        'Amodei等将奖励骇客识别为关键安全问题。'
        '在离线RL中，Levine等表明分布外动作可利用价值函数高估。'
        '本文的因果校准方法与RLHF中的奖励模型修正方法相关，'
        '但使用因果推断而非人类反馈来锚定修正。')
    add_heading_cn(doc, '2.3 DRL用于耕地整治', level=2)
    add_para(doc,
        '本工作建立在渐进式扩展计划之上：'
        '合成数据上的地块级DRL、真实地籍数据上的可扩展性分析、'
        '区块级空间抽象在单个乡镇上实现100%收敛、'
        '以及多智能体县域级协调。'
        '本文解决了县域级方法中识别的计算成本和可靠性限制。')
    add_heading_cn(doc, '2.4 因果推断用于奖励校准', level=2)
    add_para(doc,
        '倾向得分方法通过平衡处理组和对照组之间的混杂因素，'
        '从观测数据中估计处理效应。'
        '我们将此框架适配到RL设置：'
        '"处理"是选择高潜力区块，"结果"是产生的奖励，'
        '"混杂因素"是影响区块选择和奖励的全局状态特征。'
        '据我们所知，这是首次使用因果推断校准学习环境的奖励函数。')


def write_sec3(doc):
    add_heading_cn(doc, '3. 问题设定')
    add_para(doc,
        '我们考虑县域级耕地整治MDP。'
        '状态s_t包含13个乡镇2,600个整治区块的逐区块特征和全局县域指标，总维度44,212。'
        '每步智能体选择一个区块（对无效区块使用动作掩码），'
        '环境使用连通性感知贪心引擎在该区块内执行5次耕地-林地交换。'
        '回合持续100步（总预算：500次交换）。'
        '奖励结合坡度降低、连片度改善和百亩方形成指标。'
        '无模型训练需500,000个时步，每种子8-12小时A100，33%失败率。')


def write_sec4(doc):
    add_heading_cn(doc, '4. 方法')
    add_para(doc,
        '本方法包含三个阶段：(1)轨迹收集；(2)转移模型训练；(3)策略优化+可选因果校准。')
    add_heading_cn(doc, '4.1 轨迹收集', level=2)
    add_para(doc,
        '从真实环境收集12,000条状态-动作-奖励-下一状态转移，'
        '使用随机策略（6,000条）和贪心策略（6,000条）。')
    add_heading_cn(doc, '4.2 神经转移模型', level=2)
    add_para(doc,
        '237K参数的转移模型预测下一观测和奖励。'
        '架构：区块编码器+动作嵌入+全局编码器+上下文聚合（128维），'
        '三个预测头分别输出区块残差、全局残差和奖励。'
        'CPU上60分钟训练完成，验证余弦相似度0.9998。')
    add_heading_cn(doc, '4.3 策略优化', level=2)
    add_para(doc,
        'LearnedCountyEnv封装转移模型为Gymnasium环境。'
        '使用Maskable PPO训练100K步，CPU上25-57分钟。'
        '策略在学习环境训练，在真实环境评估。')
    add_heading_cn(doc, '4.4 因果奖励校准', level=2)
    add_para(doc,
        '通过倾向得分匹配估计动作质量对奖励的真实因果效应（ATT=0.049），'
        '与学习模型预测效应（0.265）对比，得到校准因子\u03b1=0.185。'
        '训练时所有奖励乘以\u03b1，修正5.4倍的系统性高估。')


def write_sec5(doc):
    add_heading_cn(doc, '5. 实验')
    add_heading_cn(doc, '5.1 主要结果', level=2)
    add_table(doc,
        '表1 方法比较。**表示p<0.01（Mann-Whitney U检验）。',
        ['方法', '种子', '坡度(%)', '显著性', '连片度', '时间', '硬件'],
        [
            ['序列贪心', '1', '-0.27', '---', '+0.008', '10s', 'CPU'],
            ['集中式DRL', '5', '-0.79\u00b10.36', '', '+0.017', '8h', 'A100'],
            ['MARL DRL', '4', '-0.84\u00b10.07', '', '+0.017', '12h', 'A100'],
            ['基于模型', '15', '-0.976\u00b10.129', '**', '+0.013', '~2h', 'CPU'],
            ['基于模型+校准', '15', '-1.102\u00b10.100', '**', '+0.011', '~2h', 'CPU'],
        ])
    add_para(doc,
        '基于模型策略（15种子）超过MARL 16%、集中式24%。'
        '因果校准后提升至-1.102%（p=0.004）。CPU约2小时，快16倍且无需GPU。')
    add_heading_cn(doc, '5.2 因果校准消融', level=2)
    add_table(doc,
        '表2 因果校准效果（15种子）。',
        ['配置', 'n', '坡度(%)', 'p值', '胜出'],
        [
            ['无校准', '15', '-0.976\u00b10.129', '---', '---'],
            ['有校准(\u03b1=0.185)', '15', '-1.102\u00b10.100', '0.004**', '10/15'],
        ])
    add_heading_cn(doc, '5.3 因果校准 vs 网格搜索', level=2)
    add_table(doc,
        '表3 \u03b1网格搜索（每值5种子）。',
        ['\u03b1', '坡度(%)', '标准差'],
        [
            ['0.100', '-1.123', '0.151'], ['0.150', '-1.120', '0.109'],
            ['0.185(因果)', '-1.171', '0.103'], ['0.200(最优)', '-1.209', '0.092'],
            ['0.300', '-1.171', '0.077'], ['0.500', '-1.043', '0.065'],
            ['0.700', '-0.991', '0.047'], ['1.000', '-0.959', '0.055'],
        ])
    add_para(doc,
        '因果\u03b1=0.185与网格最优\u03b1=0.200仅差3.1%，证明因果校准为免调参方法。')
    add_heading_cn(doc, '5.4 GeoFM消融', level=2)
    add_para(doc,
        'GeoFM嵌入将转移模型损失降低40%但策略质量下降（-0.70% vs -0.96%），'
        '揭示模型精度与策略质量的非单调关系。')


def write_sec6(doc):
    add_heading_cn(doc, '6. 讨论')
    add_para(doc,
        '基于模型方法优于无模型的三个原因：无噪声梯度、更快迭代、隐式课程。'
        '因果校准\u03b1=0.185落在最优平台[0.15,0.30]内，与网格搜索最优仅差3%，'
        '证明因果推断提供了基于经验证据的免调参校准。'
        'GeoFM消融揭示精度-效用悖论：更精确的模型可能更忠实地再现退化吸引子。')
    add_para(doc, '局限性：分布偏移、单一研究区域、全局校准因子、奖励预测噪声。'
        '状态依赖\u03b1(s)和跨区域泛化留作未来工作。')


def write_sec7(doc):
    add_heading_cn(doc, '7. 结论')
    add_list(doc, [
        '237K参数转移模型支持有效策略训练（余弦相似度0.9998）。',
        '15种子基于模型策略实现-0.976%坡度降低，超过无模型基线。',
        '因果校准提升13.0%至-1.102%（p=0.004），与网格搜索最优仅差3%。',
        '完整流程CPU约2小时，比无模型快16倍且无需GPU。',
        'GeoFM消融揭示模型精度与策略质量非单调关系。',
    ])


def write_references(doc):
    add_heading_cn(doc, '参考文献')
    refs = [
        '[Amodei et al., 2016] Concrete problems in AI safety. arXiv:1606.06565.',
        '[Hafner et al., 2023] Mastering diverse domains through world models. arXiv:2301.04104.',
        '[Lambert et al., 2022] Objective mismatch in MBRL. L4DC.',
        '[Schrittwieser et al., 2020] Mastering Atari, Go, chess and shogi. Nature 588, 604-609.',
        '[Stuart, 2010] Matching methods for causal inference. Statistical Science 25(1), 1-21.',
        '[Zheng et al., 2023] Spatial planning via DRL. Nature Computational Science 3(9), 748-762.',
        '[Zhou and Jing, 2026a] 可迁移DRL耕地空间布局优化框架. IJGIS (已投稿).',
        '[Zhou and Jing, 2026b] DRL耕地布局优化：机遇与局限. CEAG (已投稿).',
        '[Zhou and Jing, 2026c] 从地块到田块：DRL耕地整治规划方法. LUP (已投稿).',
        '[Zhou and Jing, 2026d] 从田块到县域：多智能体DRL耕地整治协调. CEUS (已投稿).',
    ]
    for ref in refs:
        add_para(doc, ref, size=9, space_after=3)


def build_cn():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    ea = rpr.makeelement(qn('w:rFonts'), {})
    ea.set(qn('w:eastAsia'), 'SimSun')
    rpr.insert(0, ea)
    write_frontmatter(doc)
    write_abstract(doc)
    write_body(doc)
    write_references(doc)
    doc.save(OUT_PATH)
    print(f'Saved: {OUT_PATH}')


if __name__ == '__main__':
    build_cn()
    print('Done!')

